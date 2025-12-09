# ==========================================================
# utils/integration_edital.py
# SynapseNext – SAAB / TJSP
# Integração estável do módulo EDITAL (modo híbrido)
# ==========================================================

from __future__ import annotations

import os
import io
import re
import json
from pathlib import Path
from datetime import datetime
from typing import Dict, Any, Optional

# Streamlit é opcional para permitir testes sem UI
try:
    import streamlit as st  # type: ignore
except Exception:  # pragma: no cover
    st = None

# -----------------------------
# Dependências opcionais de extração
# -----------------------------
try:
    import fitz  # PyMuPDF
except Exception:  # pragma: no cover
    fitz = None

try:
    import docx2txt
except Exception:  # pragma: no cover
    docx2txt = None

try:
    from docx import Document
except Exception:  # pragma: no cover
    Document = None

# -----------------------------
# OpenAI – criação tardia (opcional)
# -----------------------------
def _get_openai_client():
    api_key = None
    if st and getattr(st, "secrets", None):
        api_key = st.secrets.get("OPENAI_API_KEY")
    api_key = api_key or os.getenv("OPENAI_API_KEY") or os.getenv("openai_api_key")

    try:
        from openai import OpenAI  # openai>=1.x
    except Exception:
        return None, None

    if not api_key:
        return None, None

    try:
        return OpenAI(api_key=api_key), "gpt-4o-mini"
    except Exception:
        return None, None


# -----------------------------
# Paths institucionais (opcional)
# -----------------------------
BASE_DIR = Path(__file__).resolve().parents[1] if "__file__" in globals() else Path.cwd()
EXPORTS_DIR = BASE_DIR / "exports"
EXPORTS_DIR.mkdir(parents=True, exist_ok=True)

KB_EDITAL_DIR = (BASE_DIR / "knowledge_base" / "edital")  # tolerante: não obrigatório


# ==========================================================
# 📚 Leitura de modelos da KB (tolerante)
# ==========================================================
def ler_modelos_edital() -> str:
    textos = []
    if KB_EDITAL_DIR.exists():
        for arq in KB_EDITAL_DIR.glob("*.txt"):
            try:
                textos.append(arq.read_text(encoding="utf-8"))
            except Exception:
                pass
    return "\n\n".join(textos)


# ==========================================================
# 🔗 Integração de contexto (DFD, ETP, TR) – modo híbrido
# ==========================================================
def integrar_com_contexto(session_state: dict | None) -> dict:
    """
    Constrói um contexto integrado com o que existir na sessão:
    DFD + ETP + TR. É tolerante à ausência de qualquer um deles.
    """
    ss = session_state or {}
    contexto = {}
    for chave in ("dfd_campos_ai", "etp_campos_ai", "tr_campos_ai"):
        bloco = ss.get(chave)
        if isinstance(bloco, dict) and bloco:
            contexto[chave] = bloco
    return contexto

# Alias de compatibilidade
def consolidar_contexto(session_state: dict | None) -> dict:
    return integrar_com_contexto(session_state)


# ==========================================================
# 🧾 Extração e limpeza de texto do insumo
# ==========================================================
def extrair_texto_arquivo(arquivo) -> str:
    nome = getattr(arquivo, "name", "").lower()

    def limpar(txt: str) -> str:
        txt = re.sub(r"\s+", " ", txt or "")
        txt = re.sub(r"[^\w\s.,;:!?()/%\-–—ºª°]", "", txt)
        return txt.strip()

    try:
        # PDF
        if nome.endswith(".pdf"):
            data = arquivo.read()
            try:
                arquivo.seek(0)
            except Exception:
                pass
            if fitz is not None:
                try:
                    texto = ""
                    with fitz.open(stream=data, filetype="pdf") as pdf:
                        for p in pdf:
                            texto += p.get_text("text") + "\n"
                    return limpar(texto)
                except Exception:
                    pass
            return ""  # sem fallback obrigatório p/ PDF (evita hard deps)

        # DOCX
        if nome.endswith(".docx") and docx2txt is not None:
            data = arquivo.read()
            try:
                arquivo.seek(0)
            except Exception:
                pass
            try:
                return limpar(docx2txt.process(io.BytesIO(data)))
            except Exception:
                pass

        # TXT
        if nome.endswith(".txt"):
            data = arquivo.read()
            try:
                arquivo.seek(0)
            except Exception:
                pass
            try:
                return limpar(data.decode("utf-8", errors="ignore"))
            except Exception:
                return limpar(data.decode("latin-1", errors="ignore"))
    except Exception:
        pass
    return ""


# ==========================================================
# 🧠 IA opcional para estruturar campos do edital
# ==========================================================
def _chamar_ia_edital(texto_insumo: str, modelos: str, contexto: dict) -> Dict[str, Any]:
    client, model = _get_openai_client()
    if client is None or not model or not texto_insumo.strip():
        return {}

    system_prompt = (
        "Você é um agente institucional do TJSP (SAAB) especializado em minutas de EDITAL (Lei 14.133/2021). "
        "Use linguagem padrão SAAB/TJSP. Retorne somente JSON."
    )

    user_prompt = f"""
Texto do insumo (base):
\"\"\"{texto_insumo[:10000]}\"\"\"

Contexto cumulativo disponível (DFD, ETP, TR):
\"\"\"{json.dumps(contexto or {}, ensure_ascii=False)}\"\"\"

Modelos institucionais (KB):
\"\"\"{(modelos or '')[:8000]}\"\"\"

Retorne APENAS um JSON com os campos:
{{
  "objeto": "",
  "tipo_licitacao": "",
  "criterio_julgamento": "",
  "condicoes_participacao": "",
  "exigencias_habilitacao": "",
  "obrigacoes_contratada": "",
  "prazo_execucao": "",
  "fontes_recursos": "",
  "gestor_fiscal": "",
  "observacoes_gerais": ""
}}
"""

    try:
        resp = client.chat.completions.create(
            model=model,
            messages=[{"role": "system", "content": system_prompt},
                      {"role": "user", "content": user_prompt}],
            temperature=0.3,
            max_tokens=900,
        )
        conteudo = (resp.choices[0].message.content or "").strip()
        m = re.search(r"\{.*\}", conteudo, re.S)
        return json.loads(m.group(0)) if m else {}
    except Exception:
        return {}


# ==========================================================
# 🧹 Normalização e defaults (modo híbrido)
# ==========================================================
def _normalizar_campos(campos: Dict[str, Any], contexto: dict) -> Dict[str, str]:
    hoje = datetime.now()
    numero_auto = f"TJSP-PE-{hoje.year}-{hoje.strftime('%m')}{hoje.strftime('%d')}"
    data_auto = hoje.strftime("%d/%m/%Y")

    dfd = (contexto or {}).get("dfd_campos_ai", {})
    etp = (contexto or {}).get("etp_campos_ai", {})
    tr  = (contexto or {}).get("tr_campos_ai", {})

    def first_nonempty(*vals):
        for v in vals:
            if isinstance(v, str) and v.strip():
                return v
        return ""

    defaults = {
        "objeto": first_nonempty(tr.get("objeto"), etp.get("objeto"), dfd.get("objeto")),
        "tipo_licitacao": "Pregão eletrônico",
        "criterio_julgamento": first_nonempty(tr.get("criterios_de_julgamento"), "Menor preço global"),
        "condicoes_participacao": "",
        "exigencias_habilitacao": "",
        "obrigacoes_contratada": first_nonempty(tr.get("obrigacoes_da_contratada")),
        "prazo_execucao": first_nonempty(tr.get("prazo_execucao")),
        "fontes_recursos": first_nonempty(tr.get("fonte_recurso")),
        "gestor_fiscal": first_nonempty(dfd.get("responsavel"), dfd.get("responsavel_tecnico")),
        "observacoes_gerais": "",
        "numero_edital": numero_auto,
        "data_publicacao": data_auto,
        "unidade_solicitante": first_nonempty(dfd.get("unidade_solicitante")),
        "responsavel": first_nonempty(dfd.get("responsavel"), dfd.get("responsavel_tecnico")),
    }

    result = {}
    for k, v in defaults.items():
        result[k] = (campos.get(k) or v or "").strip()
        result[k] = re.sub(r"\s+", " ", result[k]).strip()

    return result

# Alias compatível com páginas antigas
def normalizar_campos_edital(campos: Dict[str, Any], contexto: dict) -> Dict[str, str]:
    return _normalizar_campos(campos, contexto)


# ==========================================================
# 🧱 Geração de rascunho e DOCX
# ==========================================================
def gerar_rascunho_edital(campos: Dict[str, str], modelos_referencia: str = "") -> str:
    """
    Gera rascunho textual do Edital de forma estruturada e profissional.
    Agora com formatação enriquecida e seções detalhadas.
    """
    
    # Cabeçalho oficial
    linhas = [
        "=" * 80,
        "TRIBUNAL DE JUSTIÇA DO ESTADO DE SÃO PAULO",
        "DIRETORIA EXECUTIVA DE GESTÃO DE SUPRIMENTOS",
        "=" * 80,
        "",
        f"EDITAL DE LICITAÇÃO Nº {campos.get('numero_edital', 'XXXXX/YYYY')}",
        f"PROCESSO ADMINISTRATIVO: {campos.get('numero_edital', 'XXXXX/YYYY')}",
        "",
        f"Data de Publicação: {campos.get('data_publicacao', '__/__/____')}",
        f"Unidade Solicitante: {campos.get('unidade_solicitante', 'A definir')}",
        f"Responsável: {campos.get('responsavel', 'A definir')}",
        "",
        "=" * 80,
        "",
        "O TRIBUNAL DE JUSTIÇA DO ESTADO DE SÃO PAULO, por meio da Diretoria Executiva",
        "de Gestão de Suprimentos, torna público que realizará licitação na modalidade",
        "e critério abaixo especificados, regida pela Lei Federal nº 14.133/2021 e",
        "demais normas pertinentes.",
        "",
        "=" * 80,
        "",
    ]
    
    # 1. OBJETO
    linhas.extend([
        "1. DO OBJETO",
        "-" * 80,
        "",
        campos.get("objeto", "(Descrição do objeto não informada)"),
        "",
        "1.1. A contratação será regida pela Lei Federal nº 14.133/2021 e demais",
        "     normas complementares.",
        "",
        "=" * 80,
        "",
    ])
    
    # 2. TIPO E CRITÉRIO
    linhas.extend([
        "2. DA MODALIDADE E CRITÉRIO DE JULGAMENTO",
        "-" * 80,
        "",
        f"2.1. Modalidade: {campos.get('tipo_licitacao', 'Pregão Eletrônico')}",
        "",
        f"2.2. Critério de Julgamento: {campos.get('criterio_julgamento', 'Menor Preço')}",
        "",
        "2.3. Modo de Disputa: Aberto e fechado, conforme art. 17, §2º da Lei 14.133/2021.",
        "",
        "=" * 80,
        "",
    ])
    
    # 3. CONDIÇÕES DE PARTICIPAÇÃO
    condicoes = campos.get("condicoes_participacao", "")
    linhas.extend([
        "3. DAS CONDIÇÕES DE PARTICIPAÇÃO",
        "-" * 80,
        "",
        condicoes if condicoes else (
            "3.1. Poderão participar desta licitação as empresas:\n"
            "     a) Regularmente estabelecidas no País;\n"
            "     b) Que atendam às condições de habilitação previstas neste Edital;\n"
            "     c) Credenciadas no portal de compras governamental (quando aplicável);\n"
            "     d) Que não estejam suspensas ou impedidas de licitar com a Administração Pública.\n"
            "\n"
            "3.2. Não poderão participar:\n"
            "     a) Empresas em processo de falência, recuperação judicial ou dissolução;\n"
            "     b) Empresas declaradas inidôneas pela Administração Pública;\n"
            "     c) Empresas com vínculo de parentesco/subordinação com agentes públicos do TJSP."
        ),
        "",
        "=" * 80,
        "",
    ])
    
    # 4. HABILITAÇÃO (CRÍTICO - deve ser detalhado)
    habilitacao = campos.get("exigencias_habilitacao", "")
    linhas.extend([
        "4. DA HABILITAÇÃO",
        "-" * 80,
        "",
        "Os licitantes deverão apresentar os seguintes documentos:",
        "",
        habilitacao if habilitacao else (
            "4.1. HABILITAÇÃO JURÍDICA:\n"
            "     a) Registro comercial (Junta Comercial) para empresas individuais;\n"
            "     b) Ato constitutivo, estatuto ou contrato social em vigor;\n"
            "     c) Decreto de autorização (empresas estrangeiras);\n"
            "     d) Cartão CNPJ atualizado.\n"
            "\n"
            "4.2. REGULARIDADE FISCAL E TRABALHISTA:\n"
            "     a) Certidão Negativa de Débitos Federais (Receita Federal/PGFN);\n"
            "     b) Certidão Negativa Estadual;\n"
            "     c) Certidão Negativa Municipal (sede do licitante);\n"
            "     d) Certidão Negativa de Débitos Trabalhistas (TST);\n"
            "     e) Certidão de Regularidade FGTS;\n"
            "     f) Certidão Negativa de Débitos INSS.\n"
            "\n"
            "4.3. QUALIFICAÇÃO TÉCNICA:\n"
            "     a) Atestado(s) de capacidade técnica emitido(s) por pessoa jurídica de\n"
            "        direito público ou privado, comprovando execução de serviços compatíveis;\n"
            "     b) Registro ou inscrição na entidade profissional competente (quando aplicável);\n"
            "     c) Certidão de Acervo Técnico (CAT) dos responsáveis técnicos.\n"
            "\n"
            "4.4. QUALIFICAÇÃO ECONÔMICO-FINANCEIRA:\n"
            "     a) Balanço patrimonial do último exercício social;\n"
            "     b) Certidão negativa de falência ou recuperação judicial (90 dias validade);\n"
            "     c) Patrimônio líquido mínimo de 10% do valor estimado da contratação."
        ),
        "",
        "=" * 80,
        "",
    ])
    
    # 5. OBRIGAÇÕES DA CONTRATADA (CRÍTICO)
    obrigacoes = campos.get("obrigacoes_contratada", "")
    linhas.extend([
        "5. DAS OBRIGAÇÕES DA CONTRATADA",
        "-" * 80,
        "",
        obrigacoes if obrigacoes else (
            "5.1. Executar os serviços em conformidade com as especificações técnicas;\n"
            "\n"
            "5.2. Responsabilizar-se por todos os encargos trabalhistas, previdenciários,\n"
            "     fiscais e comerciais resultantes da execução do contrato;\n"
            "\n"
            "5.3. Manter durante toda a execução do contrato as condições de habilitação;\n"
            "\n"
            "5.4. Fornecer mão de obra qualificada e devidamente treinada;\n"
            "\n"
            "5.5. Apresentar relatórios mensais de atividades executadas;\n"
            "\n"
            "5.6. Reparar ou corrigir, às suas expensas, vícios, defeitos ou incorreções;\n"
            "\n"
            "5.7. Manter sigilo sobre informações obtidas durante a execução dos serviços."
        ),
        "",
        "=" * 80,
        "",
    ])
    
    # 6. PRAZO
    linhas.extend([
        "6. DO PRAZO DE EXECUÇÃO",
        "-" * 80,
        "",
        f"6.1. Prazo de vigência: {campos.get('prazo_execucao', 'A definir')}",
        "",
        "6.2. O prazo poderá ser prorrogado mediante termo aditivo, desde que atendidos",
        "     os requisitos do art. 107 da Lei 14.133/2021.",
        "",
        "=" * 80,
        "",
    ])
    
    # 7. RECURSOS
    linhas.extend([
        "7. DOS RECURSOS ORÇAMENTÁRIOS",
        "-" * 80,
        "",
        f"7.1. Fonte de Recursos: {campos.get('fontes_recursos', 'A definir')}",
        "",
        "7.2. As despesas decorrentes da presente contratação correrão à conta dos",
        "     recursos consignados no orçamento do TJSP.",
        "",
        "=" * 80,
        "",
    ])
    
    # 8. GESTOR/FISCAL
    linhas.extend([
        "8. DA GESTÃO E FISCALIZAÇÃO DO CONTRATO",
        "-" * 80,
        "",
        f"8.1. Gestor do Contrato: {campos.get('gestor_fiscal', 'A definir')}",
        "",
        "8.2. A fiscalização do contrato será exercida nos termos do art. 117 da",
        "     Lei 14.133/2021, cabendo ao fiscal acompanhar e atestar a execução dos serviços.",
        "",
        "=" * 80,
        "",
    ])
    
    # 9. DISPOSIÇÕES FINAIS
    obs = campos.get("observacoes_gerais", "")
    linhas.extend([
        "9. DAS DISPOSIÇÕES FINAIS",
        "-" * 80,
        "",
        obs if obs else (
            "9.1. Eventuais dúvidas deverão ser encaminhadas por escrito à Comissão de Licitação.\n"
            "\n"
            "9.2. O Edital completo encontra-se disponível no Portal de Compras do TJSP.\n"
            "\n"
            "9.3. A Administração poderá revogar a licitação por razões de interesse público\n"
            "     ou anulá-la por ilegalidade, assegurado o contraditório e ampla defesa.\n"
            "\n"
            "9.4. Os casos omissos serão decididos pela Comissão de Licitação, com base\n"
            "     na Lei 14.133/2021 e demais normas aplicáveis."
        ),
        "",
        "=" * 80,
        "",
        f"São Paulo, {campos.get('data_publicacao', '__/__/____')}",
        "",
        "",
        "_" * 60,
        "Presidente da Comissão de Licitação",
        "Tribunal de Justiça do Estado de São Paulo",
        "",
    ])
    
    return "\n".join(linhas)
    if modelos_referencia:
        linhas += ["", "ANEXO – ORIENTAÇÕES INSTITUCIONAIS (KB)", (modelos_referencia[:1200] + " …")]
    return "\n".join(linhas)


def gerar_edital_docx(campos: Dict[str, str], texto_completo: Optional[str] = None, session_state: dict = None) -> Optional[str]:
    """
    Gera documento DOCX do edital com formatação profissional.
    
    Args:
        campos: dicionário com os 12 campos do edital
        texto_completo: rascunho textual completo (opcional)
        session_state: dict do streamlit session_state para salvar buffer
    
    Returns:
        caminho do arquivo se salvo em disco, None se apenas em buffer
    """
    if Document is None:
        return None  # evita falha se python-docx não estiver instalado

    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    
    # Cabeçalho institucional
    header = doc.add_heading('TRIBUNAL DE JUSTIÇA DO ESTADO DE SÃO PAULO', level=1)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subheader = doc.add_paragraph('DIRETORIA EXECUTIVA DE GESTÃO DE SUPRIMENTOS')
    subheader.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subheader.runs[0].bold = True
    
    doc.add_paragraph()  # espaço
    
    # Título do Edital
    titulo = doc.add_heading(f"EDITAL DE LICITAÇÃO Nº {campos.get('numero_edital', 'XXXXX/YYYY')}", level=1)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadados
    meta = doc.add_paragraph()
    meta.add_run(f"Data de Publicação: ").bold = True
    meta.add_run(campos.get('data_publicacao', '__/__/____'))
    
    meta = doc.add_paragraph()
    meta.add_run(f"Unidade Solicitante: ").bold = True
    meta.add_run(campos.get('unidade_solicitante', 'A definir'))
    
    meta = doc.add_paragraph()
    meta.add_run(f"Responsável: ").bold = True
    meta.add_run(campos.get('responsavel', 'A definir'))
    
    doc.add_paragraph()  # espaço
    
    # Preâmbulo
    preambulo = doc.add_paragraph(
        "O TRIBUNAL DE JUSTIÇA DO ESTADO DE SÃO PAULO, por meio da Diretoria Executiva "
        "de Gestão de Suprimentos, torna público que realizará licitação na modalidade "
        "e critério abaixo especificados, regida pela Lei Federal nº 14.133/2021 e "
        "demais normas pertinentes."
    )
    preambulo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph()  # espaço

    def bloco_estruturado(titulo: str, conteudo: str):
        """Adiciona seção com formatação profissional."""
        h = doc.add_heading(titulo, level=2)
        h.runs[0].font.color.rgb = RGBColor(0, 51, 102)  # Azul institucional
        
        if conteudo and conteudo.strip():
            # Se tiver quebras de linha, preservar formatação
            paragrafos = conteudo.split('\n\n')
            for p_text in paragrafos:
                if p_text.strip():
                    p = doc.add_paragraph(p_text.strip())
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        else:
            p = doc.add_paragraph("(Informação não disponível)")
            p.runs[0].italic = True
        
        doc.add_paragraph()  # espaço entre seções
    
    # Seções do Edital
    bloco_estruturado(
        "1. DO OBJETO",
        campos.get("objeto", "")
    )
    
    bloco_estruturado(
        "2. DA MODALIDADE E CRITÉRIO DE JULGAMENTO",
        f"Modalidade: {campos.get('tipo_licitacao', 'Pregão Eletrônico')}\n\n"
        f"Critério de Julgamento: {campos.get('criterio_julgamento', 'Menor Preço')}\n\n"
        f"Modo de Disputa: Aberto e fechado, conforme art. 17, §2º da Lei 14.133/2021."
    )
    
    bloco_estruturado(
        "3. DAS CONDIÇÕES DE PARTICIPAÇÃO",
        campos.get("condicoes_participacao", "")
    )
    
    bloco_estruturado(
        "4. DA HABILITAÇÃO",
        campos.get("exigencias_habilitacao", "")
    )
    
    bloco_estruturado(
        "5. DAS OBRIGAÇÕES DA CONTRATADA",
        campos.get("obrigacoes_contratada", "")
    )
    
    bloco_estruturado(
        "6. DO PRAZO DE EXECUÇÃO",
        campos.get("prazo_execucao", "")
    )
    
    bloco_estruturado(
        "7. DOS RECURSOS ORÇAMENTÁRIOS",
        campos.get("fontes_recursos", "")
    )
    
    bloco_estruturado(
        "8. DA GESTÃO E FISCALIZAÇÃO DO CONTRATO",
        campos.get("gestor_fiscal", "")
    )
    
    bloco_estruturado(
        "9. DAS DISPOSIÇÕES FINAIS",
        campos.get("observacoes_gerais", "")
    )

    # Anexo com rascunho completo (se fornecido)
    if texto_completo:
        doc.add_page_break()
        anexo_titulo = doc.add_heading("ANEXO – RASCUNHO INTEGRAL", level=1)
        anexo_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Adicionar texto completo preservando formatação
        for linha in texto_completo.split("\n"):
            if linha.strip():
                if linha.startswith("=") or linha.startswith("-"):
                    continue  # pular separadores decorativos
                p = doc.add_paragraph(linha)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Rodapé
    doc.add_paragraph()
    rodape = doc.add_paragraph(f"São Paulo, {campos.get('data_publicacao', '__/__/____')}")
    rodape.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    assinatura = doc.add_paragraph("_" * 60)
    assinatura.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    cargo = doc.add_paragraph("Presidente da Comissão de Licitação")
    cargo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    instituicao = doc.add_paragraph("Tribunal de Justiça do Estado de São Paulo")
    instituicao.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Garantir que o diretório existe antes de salvar
    nome_arquivo = f"Edital_{campos.get('numero_edital','TJSP-PE').replace('/', '-')}.docx"
    
    # Usar diretório exports que já está garantido
    exports_path = Path(EXPORTS_DIR)
    exports_path.mkdir(parents=True, exist_ok=True)
    
    caminho = str(exports_path / nome_arquivo)
    
    # SEMPRE criar o buffer primeiro (para garantir que funciona)
    buffer = io.BytesIO()
    try:
        doc.save(buffer)
        buffer.seek(0)
        print(f"[integration_edital] Buffer DOCX criado: {len(buffer.getvalue())} bytes")
    except Exception as e:
        print(f"[integration_edital] ERRO CRÍTICO ao criar buffer: {e}")
        import traceback
        traceback.print_exc()
        return None
    
    # Salvar o buffer no session_state SEMPRE (prioridade)
    if session_state is not None:
        try:
            session_state["edital_docx_buffer"] = buffer
            session_state["edital_docx_nome"] = nome_arquivo
            print(f"[integration_edital] ✅ Buffer salvo no session_state: {nome_arquivo}")
        except Exception as e:
            print(f"[integration_edital] ERRO ao salvar no session_state: {e}")
            import traceback
            traceback.print_exc()
    else:
        print(f"[integration_edital] ⚠️ session_state é None, não foi possível salvar buffer")
    
    # Tentar salvar no disco (opcional, só funciona em Codespaces)
    try:
        # Criar novo buffer para o arquivo (o anterior já está no session_state)
        doc.save(caminho)
        print(f"[integration_edital] DOCX salvo em disco: {caminho}")
        return caminho
    except Exception as e:
        print(f"[integration_edital] Não foi possível salvar em disco (esperado no Streamlit Cloud): {e}")
        # Retornar None mas buffer já está no session_state
        return None


# ==========================================================
# 🚀 Função principal (entrada unificada para a página)
# ==========================================================
def processar_insumo_edital(arquivo, contexto_previo: dict | None = None, artefato: str = "EDITAL") -> dict:
    """
    1) Extrai texto do insumo (PDF/DOCX/TXT).
    2) Integra contexto (DFD/ETP/TR) – modo híbrido.
    3) Chama IA (se disponível) para estruturar campos.
    4) Normaliza e gera rascunho + DOCX.
    5) Retorna dicionário padronizado para pré-preenchimento.
    """
    texto = extrair_texto_arquivo(arquivo)
    if not texto:
        return {"erro": "Falha na extração de texto do insumo de EDITAL."}

    contexto = contexto_previo or integrar_com_contexto(st.session_state if st else None)
    modelos = ler_modelos_edital()
    campos_ia = _chamar_ia_edital(texto, modelos, contexto)
    campos = _normalizar_campos(campos_ia if isinstance(campos_ia, dict) else {}, contexto)

    rascunho = gerar_rascunho_edital(campos, modelos_referencia="")
    docx_path = gerar_edital_docx(campos, texto_completo=rascunho)

    payload = {
        "artefato": artefato,
        "nome_arquivo": getattr(arquivo, "name", ""),
        "status": "processado",
        "campos_ai": campos,
        "docx_path": docx_path,
        "contexto_usado": list((contexto or {}).keys()),
    }

    if st is not None:
        st.session_state["last_edital"] = payload

    return payload


# ==========================================================
# 🔁 Wrappers de compatibilidade (páginas antigas)
# ==========================================================
def processar_edital_dinamico(arquivo, contexto_previo: dict | None = None, artefato: str = "EDITAL") -> dict:
    """Alias histórico usado por algumas páginas."""
    return processar_insumo_edital(arquivo, contexto_previo=contexto_previo, artefato=artefato)


# ==========================================================
# 🤖 Geração de Edital com IA (integração com EditalAgent)
# ==========================================================
def gerar_edital_com_ia(contexto_previo: dict = None) -> dict:
    """
    Carrega insumo bruto do INSUMOS e processa com EditalAgent.
    Salva resultado estruturado em exports/edital_data.json
    
    Args:
        contexto_previo: dict com dados de DFD/ETP/TR (opcional)
    
    Returns:
        dict com estrutura Edital completa (12 campos)
    """
    from agents.edital_agent import processar_edital_com_ia
    
    # Carregar insumo bruto do módulo INSUMOS
    INSUMO_EDITAL_PATH = EXPORTS_DIR / "insumos" / "json" / "EDITAL_ultimo.json"
    
    if not INSUMO_EDITAL_PATH.exists():
        return {"erro": "Nenhum insumo EDITAL encontrado. Faça upload no módulo INSUMOS primeiro."}
    
    try:
        with open(INSUMO_EDITAL_PATH, "r", encoding="utf-8") as f:
            insumo_data = json.load(f)
    except Exception as e:
        return {"erro": f"Erro ao ler insumo: {e}"}
    
    # Obter texto bruto
    conteudo_textual = insumo_data.get("conteudo_textual", "")
    if not conteudo_textual or len(conteudo_textual) < 50:
        return {"erro": "Insumo carregado não possui texto suficiente para processamento."}
    
    # Processar com EditalAgent (com contexto DFD/ETP/TR se disponível)
    resultado_ia = processar_edital_com_ia(conteudo_textual, contexto_previo)
    
    if "erro" in resultado_ia:
        return resultado_ia
    
    # Estrutura final com metadados
    edital_processado = resultado_ia.get("EDITAL", {})
    
    dados_completos = {
        "artefato": "EDITAL",
        "timestamp": datetime.now().isoformat(),
        "arquivo_original": insumo_data.get("arquivo_original", ""),
        "data_processamento": insumo_data.get("data_processamento", ""),
        "texto_completo": conteudo_textual,
        "processado_ia": True,
        "timestamp_ia": datetime.now().isoformat(),
        "EDITAL": edital_processado,
        "contexto_usado": list((contexto_previo or {}).keys())
    }
    
    # Salvar resultado estruturado
    EDITAL_JSON_PATH = EXPORTS_DIR / "edital_data.json"
    EDITAL_JSON_PATH.parent.mkdir(parents=True, exist_ok=True)
    
    with open(EDITAL_JSON_PATH, "w", encoding="utf-8") as f:
        json.dump(dados_completos, f, ensure_ascii=False, indent=2)
    
    # Gerar também o DOCX (passando session_state explicitamente)
    print(f"[gerar_edital_com_ia] Iniciando geração do DOCX...")
    print(f"[gerar_edital_com_ia] st is None: {st is None}")
    
    # Obter session_state se disponível
    session_state_param = None
    if st is not None:
        try:
            session_state_param = st.session_state
            print(f"[gerar_edital_com_ia] session_state obtido, keys antes: {list(session_state_param.keys())}")
        except Exception as e:
            print(f"[gerar_edital_com_ia] Erro ao obter session_state: {e}")
    
    rascunho = gerar_rascunho_edital(edital_processado)
    print(f"[gerar_edital_com_ia] Rascunho gerado: {len(rascunho)} chars")
    
    docx_path = gerar_edital_docx(edital_processado, texto_completo=rascunho, session_state=session_state_param)
    print(f"[gerar_edital_com_ia] gerar_edital_docx() retornou: {docx_path}")
    
    dados_completos["docx_path"] = docx_path
    
    # CRITICAL: Verificar se buffer foi criado no session_state
    if session_state_param is not None and "edital_docx_buffer" in session_state_param:
        dados_completos["docx_buffer_disponivel"] = True
        buffer_size = len(session_state_param["edital_docx_buffer"].getvalue())
        print(f"[gerar_edital_com_ia] ✅ Buffer confirmado no session_state: {buffer_size} bytes")
        print(f"[gerar_edital_com_ia] session_state keys depois: {list(session_state_param.keys())}")
    else:
        dados_completos["docx_buffer_disponivel"] = False
        print(f"[gerar_edital_com_ia] ❌ Buffer NÃO encontrado no session_state")
        if session_state_param is not None:
            print(f"[gerar_edital_com_ia]    session_state keys: {list(session_state_param.keys())}")
        print(f"[gerar_edital_com_ia] ✅ Buffer confirmado no session_state: {buffer_size} bytes")
    else:
        dados_completos["docx_buffer_disponivel"] = False
        print(f"[gerar_edital_com_ia] ❌ Buffer NÃO encontrado no session_state")
        print(f"[gerar_edital_com_ia]    st is None: {st is None}")
        if st is not None:
            print(f"[gerar_edital_com_ia]    session_state keys: {list(st.session_state.keys())}")
    
    return dados_completos
