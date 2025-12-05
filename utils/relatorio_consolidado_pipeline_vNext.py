# -*- coding: utf-8 -*-
"""
🔧 relatorio_consolidado_pipeline_vNext.py
==============================================================
Pipeline consolidado para auditoria técnica e geração de
relatórios institucionais no ecossistema SynapseNext vNext+.

Inclui:
- Coleta de snapshots dos artefatos (DFD, ETP, TR, Edital)
- Validação semântica assistida por IA
- Comparação interdocumental (Coerência Global)
- Normalização dos dados para o Relatório Técnico

Autor: Equipe Synapse.Engineer
Instituição: Secretaria de Administração e Abastecimento – TJSP
Versão: vNext+ (SAAB 5.0)
==============================================================
"""

import os
import json
from datetime import datetime
import pandas as pd

# Dependências internas
try:
    from utils.analytics_engine_vNext import analisar_coerencia_global
    from knowledge.validators.validator_engine_vNext import validar_semantica_ia
    from utils.comparador_pipeline import comparar_documentos
    from utils.export_snapshot import carregar_snapshot
except ImportError:
    pass


# ==========================================================
# 🔹 Função: coletar_dados_relatorio
# ==========================================================
def coletar_dados_relatorio():
    """
    Compila dados de auditoria, coerência e validação IA
    a partir dos artefatos disponíveis nos snapshots exportados.

    Retorna:
        dict {
            "ordem": [DFD, ETP, TR, Edital],
            "validacoes": { artefato: { pontuacao, mensagens } },
            "coerencia": { coerencia_global, divergencias, ausencias },
            "timestamp": datetime,
        }
    """
    dados = {
        "ordem": ["DFD", "ETP", "TR", "Edital"],
        "validacoes": {},
        "coerencia": {},
        "timestamp": datetime.now().isoformat()
    }

    artefatos = ["DFD", "ETP", "TR", "Edital"]
    snapshots_dir = os.path.join("exports", "snapshots")

    if not os.path.exists(snapshots_dir):
        raise FileNotFoundError("Diretório de snapshots não encontrado.")

    # Loop principal: ler e validar cada artefato
    for nome in artefatos:
        arquivo = os.path.join(snapshots_dir, f"{nome}_snapshot.json")
        if not os.path.exists(arquivo):
            dados["validacoes"][nome] = {"pontuacao": 0, "mensagens": ["Snapshot ausente."]}
            continue

        with open(arquivo, "r", encoding="utf-8") as f:
            conteudo = json.load(f)

        # Validação semântica IA
        try:
            resultado_val = validar_semantica_ia(conteudo)
            pontuacao = resultado_val.get("pontuacao", 0)
            mensagens = resultado_val.get("mensagens", [])
        except Exception as e:
            pontuacao = 0
            mensagens = [f"Erro durante a validação IA: {e}"]

        dados["validacoes"][nome] = {
            "pontuacao": pontuacao,
            "mensagens": mensagens
        }

    # ======================================================
    # Comparação interdocumental (Coerência Global)
    # ======================================================
    try:
        coerencia = analisar_coerencia_global(dados["validacoes"])
        dados["coerencia"] = coerencia
    except Exception:
        try:
            coerencia = comparar_documentos(dados["validacoes"])
            dados["coerencia"] = coerencia
        except Exception as e:
            dados["coerencia"] = {"coerencia_global": 0, "erro": str(e)}

    return dados


# ==========================================================
# 🔹 Função: gerar_relatorio_docx
# ==========================================================
def gerar_relatorio_docx(dados: dict) -> str:
    """
    Gera um arquivo .docx institucional consolidando as análises
    técnicas e semânticas executadas.

    Parâmetros:
        dados (dict): dados retornados de coletar_dados_relatorio()

    Retorna:
        str: caminho absoluto do arquivo gerado (.docx)
    """
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading("RELATÓRIO TÉCNICO CONSOLIDADO", level=1)
    doc.add_paragraph("Sistema SynapseNext vNext+ • SAAB/TJSP")
    doc.add_paragraph(f"Gerado em: {datetime.now():%d/%m/%Y %H:%M}")

    doc.add_heading("1. Sumário Executivo", level=2)
    doc.add_paragraph(
        "Este relatório consolida os resultados da auditoria digital, "
        "validação semântica e comparação interdocumental dos artefatos "
        "institucionais (DFD, ETP, TR e Edital)."
    )

    # Coerência Global
    coe = dados.get("coerencia", {})
    coerencia_global = coe.get("coerencia_global", 0)
    doc.add_heading("2. Coerência Global", level=2)
    doc.add_paragraph(f"Nível de Coerência Global: {coerencia_global:.2f}%")

    if coe.get("divergencias"):
        doc.add_heading("Divergências Identificadas", level=3)
        for d in coe["divergencias"]:
            doc.add_paragraph(f"- {d.get('descricao', '')}")

    if coe.get("ausencias"):
        doc.add_heading("Ausências Relevantes", level=3)
        for a in coe["ausencias"]:
            doc.add_paragraph(f"- {a.get('descricao', '')}")

    # Validações por artefato
    doc.add_heading("3. Validação Semântica IA por Artefato", level=2)
    for nome, v in dados.get("validacoes", {}).items():
        doc.add_paragraph(f"{nome}: {v.get('pontuacao', 0)}%")
        for msg in v.get("mensagens", []):
            doc.add_paragraph(f"  - {msg}", style="List Bullet")

    # Rodapé
    doc.add_paragraph("---")
    doc.add_paragraph(
        "Relatório técnico institucional automatizado – SynapseNext vNext+ / SAAB 5.0",
    )

    output_dir = os.path.join("exports", "relatorios")
    os.makedirs(output_dir, exist_ok=True)
    out_file = os.path.join(output_dir, f"Relatorio_Tecnico_{datetime.now():%Y%m%d_%H%M}.docx")
    doc.save(out_file)
    return out_file
