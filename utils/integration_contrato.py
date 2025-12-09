# -*- coding: utf-8 -*-
# ==========================================================
# utils/integration_contrato.py – SynapseNext / SAAB TJSP v2025.1
# ==========================================================
# - Processa insumos (PDF/DOCX/TXT) para CONTRATO com ContratoAgent
# - Normaliza campos para o formulário do módulo Contrato
# - Exporta/Carrega JSON em exports/contrato_data.json
# - Permite fusão de contexto com DFD/ETP/TR/Edital
# ==========================================================

import os
import re
import json
import io
from io import BytesIO
from typing import Dict, Any, Optional
from pathlib import Path
from datetime import datetime

# Dependências opcionais de extração
try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

try:
    import docx2txt
except ImportError:
    docx2txt = None

try:
    from docx import Document
except ImportError:
    Document = None

# Importar ContratoAgent
from agents.contrato_agent import processar_contrato_com_ia

# -----------------------------
# 📂 Export paths
# -----------------------------
EXPORTS_DIR = Path(__file__).resolve().parents[1] / "exports"
CONTRATO_JSON_PATH = EXPORTS_DIR / "contrato_data.json"

# -----------------------------
# 🧰 Utilitários Export/Load
# -----------------------------
def ensure_exports_dir():
    EXPORTS_DIR.mkdir(parents=True, exist_ok=True)

def export_contrato_to_json(data: Dict[str, Any], path: Path = CONTRATO_JSON_PATH) -> str:
    ensure_exports_dir()
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    print(f"[integration_contrato] JSON salvo: {path}")
    return str(path)

def load_contrato_from_json(path: Path = CONTRATO_JSON_PATH) -> Dict[str, Any]:
    try:
        if path.exists():
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f)
                print(f"[integration_contrato] JSON carregado: {path}")
                return data
    except Exception as e:
        print(f"[integration_contrato] Erro ao carregar JSON: {e}")
    return {}

# -----------------------------
# 🧾 Extração de texto
# -----------------------------
def extrair_texto_arquivo(arquivo) -> str:
    """Extrai texto de PDF, DOCX ou TXT."""
    nome = getattr(arquivo, "name", "").lower()
    
    try:
        if nome.endswith(".pdf"):
            if fitz is None:
                return ""
            dados = arquivo.read()
            arquivo.seek(0)
            texto = ""
            with fitz.open(stream=dados, filetype="pdf") as pdf:
                for p in pdf:
                    texto += p.get_text("text") + "\n"
            return re.sub(r"\s+", " ", texto).strip()

        elif nome.endswith(".docx"):
            if docx2txt is None:
                return ""
            dados = arquivo.read()
            arquivo.seek(0)
            return re.sub(r"\s+", " ", docx2txt.process(BytesIO(dados))).strip()

        elif nome.endswith(".txt"):
            dados = arquivo.read()
            arquivo.seek(0)
            return re.sub(r"\s+", " ", dados.decode("utf-8", errors="ignore")).strip()
    
    except Exception as e:
        print(f"[integration_contrato] ERRO na extração: {e}")
    
    return ""

# -----------------------------
# 🔗 Fusão de contexto cumulativo
# -----------------------------
def integrar_com_contexto(session_state: Dict[str, Any]) -> Dict[str, Any]:
    """Mescla DFD + ETP + TR + Edital para enriquecer CONTRATO."""
    contexto = {}
    for chave in ["dfd_campos_ai", "etp_campos_ai", "tr_campos_ai", "edital_campos_ai"]:
        bloco = session_state.get(chave)
        if isinstance(bloco, dict):
            contexto[chave] = bloco
    
    print(f"[integration_contrato] Contexto integrado: {list(contexto.keys())}")
    return contexto


# -----------------------------
# 🤖 Processamento IA – CONTRATO
# -----------------------------
def processar_insumo_contrato(
    arquivo,
    artefato: str = "CONTRATO",
    contexto_previo: Dict[str, Any] | None = None,
) -> Dict[str, Any]:
    """
    Processa insumo de CONTRATO com ContratoAgent.
    Retorna dict estruturado com 20 campos.
    """
    print(f"[integration_contrato] Processando arquivo: {getattr(arquivo, 'name', 'N/A')}")
    
    # Extrair texto
    texto = extrair_texto_arquivo(arquivo)
    if not texto or len(texto) < 50:
        return {"erro": "Falha na extração ou texto insuficiente do insumo de CONTRATO."}
    
    print(f"[integration_contrato] Texto extraído: {len(texto)} caracteres")
    
    # Processar com ContratoAgent
    try:
        resultado = processar_contrato_com_ia(texto, contexto_previo)
        
        if "erro" in resultado:
            return resultado
        
        # Estruturar resposta
        contrato_campos = resultado.get("CONTRATO", {})
        
        print(f"[integration_contrato] ✅ Processado: {len(contrato_campos)} campos")
        
        return {
            "artefato": artefato,
            "nome_arquivo": getattr(arquivo, "name", ""),
            "status": "processado",
            "timestamp": resultado.get("timestamp", datetime.now().isoformat()),
            "CONTRATO": contrato_campos,
        }
    
    except Exception as e:
        print(f"[integration_contrato] ERRO no processamento: {e}")
        import traceback
        traceback.print_exc()
        return {"erro": f"Falha ao processar CONTRATO: {e}"}


# -----------------------------
# 🤖 Processamento com contexto integrado (wrapper para UI)
# -----------------------------
def gerar_contrato_com_ia(contexto_previo: dict = None) -> dict:
    """
    Gera contrato usando APENAS contexto (sem insumo).
    Útil quando o usuário já tem DFD/ETP/TR/Edital completos.
    """
    print("[integration_contrato] Gerando contrato a partir do contexto...")
    
    if not contexto_previo:
        return {"erro": "Contexto vazio. Processe DFD/ETP/TR/Edital primeiro."}
    
    # Construir texto sintético do contexto
    texto_contexto = _construir_texto_do_contexto(contexto_previo)
    
    # Processar com ContratoAgent
    try:
        resultado = processar_contrato_com_ia(texto_contexto, contexto_previo)
        
        if "erro" in resultado:
            return resultado
        
        contrato_campos = resultado.get("CONTRATO", {})
        
        print(f"[integration_contrato] ✅ Contrato gerado: {len(contrato_campos)} campos")
        
        return {
            "artefato": "CONTRATO",
            "nome_arquivo": "gerado_de_contexto",
            "status": "processado",
            "timestamp": resultado.get("timestamp", datetime.now().isoformat()),
            "CONTRATO": contrato_campos,
        }
    
    except Exception as e:
        print(f"[integration_contrato] ERRO na geração: {e}")
        import traceback
        traceback.print_exc()
        return {"erro": f"Falha ao gerar CONTRATO: {e}"}


def _construir_texto_do_contexto(contexto: dict) -> str:
    """Constrói texto sintético a partir do contexto DFD/ETP/TR/Edital."""
    partes = []
    
    # DFD
    dfd = contexto.get("dfd_campos_ai", {})
    if dfd:
        partes.append(f"OBJETO: {dfd.get('objeto', '')}")
        partes.append(f"JUSTIFICATIVA: {dfd.get('justificativa', '')}")
        partes.append(f"VALOR ESTIMADO: {dfd.get('valor_estimado', '')}")
    
    # ETP
    etp = contexto.get("etp_campos_ai", {})
    if etp:
        partes.append(f"PRAZO ESTIMADO: {etp.get('prazo_estimado', '')}")
        partes.append(f"RESULTADOS: {etp.get('resultados_pretendidos', '')}")
    
    # TR
    tr = contexto.get("tr_campos_ai", {})
    if tr:
        partes.append(f"ESPECIFICAÇÃO TÉCNICA: {tr.get('especificacao_tecnica', '')}")
        partes.append(f"PRAZO DE EXECUÇÃO: {tr.get('prazo_execucao', '')}")
        partes.append(f"FONTE DE RECURSOS: {tr.get('fonte_recurso', '')}")
    
    # Edital
    edital = contexto.get("edital_campos_ai", {})
    if edital:
        partes.append(f"EDITAL Nº: {edital.get('numero_edital', '')}")
        partes.append(f"MODALIDADE: {edital.get('tipo_licitacao', '')}")
        partes.append(f"OBRIGAÇÕES DA CONTRATADA: {edital.get('obrigacoes_contratada', '')}")
    
    return "\n\n".join(p for p in partes if p)


# -----------------------------
# 📄 Geração de DOCX
# -----------------------------
def gerar_contrato_docx(campos: Dict[str, str], texto_completo: Optional[str] = None, session_state: dict = None) -> Optional[str]:
    """
    Gera documento DOCX do contrato com formatação profissional.
    
    Args:
        campos: dicionário com os 20 campos do contrato
        texto_completo: rascunho textual completo (opcional)
        session_state: dict do streamlit session_state para salvar buffer
    
    Returns:
        caminho do arquivo se salvo em disco, None se apenas em buffer
    """
    if Document is None:
        print("[integration_contrato] python-docx não disponível")
        return None

    try:
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("[integration_contrato] Módulos de formatação não disponíveis")
        # Usar versão simples sem formatação avançada
        return _gerar_contrato_docx_simples(campos, texto_completo, session_state)
    
    doc = Document()
    
    # Cabeçalho institucional
    header = doc.add_heading('TRIBUNAL DE JUSTIÇA DO ESTADO DE SÃO PAULO', level=1)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Título do Contrato
    titulo = doc.add_heading(f"CONTRATO ADMINISTRATIVO Nº {campos.get('numero_contrato', 'XXXXX/YYYY')}", level=1)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadados
    meta = doc.add_paragraph()
    meta.add_run(f"Data de Assinatura: ").bold = True
    meta.add_run(campos.get('data_assinatura', '__/__/____'))
    
    doc.add_paragraph()
    
    # Preâmbulo
    preambulo = doc.add_paragraph(
        f"O TRIBUNAL DE JUSTIÇA DO ESTADO DE SÃO PAULO, doravante denominado CONTRATANTE, "
        f"e {campos.get('partes_contratada', 'CONTRATADA')}, doravante denominada CONTRATADA, "
        f"firmam o presente Contrato Administrativo, regido pela Lei Federal nº 14.133/2021 e "
        f"demais normas pertinentes, mediante as cláusulas e condições seguintes:"
    )
    preambulo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph()

    def bloco_estruturado(titulo: str, conteudo: str):
        """Adiciona cláusula com formatação profissional."""
        h = doc.add_heading(titulo, level=2)
        h.runs[0].font.color.rgb = RGBColor(0, 51, 102)  # Azul institucional
        
        if conteudo and conteudo.strip():
            paragrafos = conteudo.split('\n\n')
            for p_text in paragrafos:
                if p_text.strip():
                    p = doc.add_paragraph(p_text.strip())
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        else:
            p = doc.add_paragraph("(Não especificado)")
            p.runs[0].italic = True
        
        doc.add_paragraph()
    
    # Cláusulas do Contrato
    bloco_estruturado("CLÁUSULA PRIMEIRA – DO OBJETO", campos.get("objeto", ""))
    bloco_estruturado("CLÁUSULA SEGUNDA – DA FUNDAMENTAÇÃO LEGAL", campos.get("fundamentacao_legal", ""))
    bloco_estruturado("CLÁUSULA TERCEIRA – DA VIGÊNCIA", campos.get("vigencia", ""))
    bloco_estruturado("CLÁUSULA QUARTA – DO VALOR GLOBAL", campos.get("valor_global", ""))
    bloco_estruturado("CLÁUSULA QUINTA – DA FORMA DE PAGAMENTO", campos.get("forma_pagamento", ""))
    bloco_estruturado("CLÁUSULA SEXTA – DO REAJUSTE", campos.get("reajuste", ""))
    bloco_estruturado("CLÁUSULA SÉTIMA – DA GARANTIA CONTRATUAL", campos.get("garantia_contratual", ""))
    bloco_estruturado("CLÁUSULA OITAVA – DAS OBRIGAÇÕES DA CONTRATADA", campos.get("obrigacoes_contratada", ""))
    bloco_estruturado("CLÁUSULA NONA – DAS OBRIGAÇÕES DA CONTRATANTE", campos.get("obrigacoes_contratante", ""))
    bloco_estruturado("CLÁUSULA DÉCIMA – DA FISCALIZAÇÃO", campos.get("fiscalizacao", ""))
    bloco_estruturado("CLÁUSULA DÉCIMA PRIMEIRA – DAS PENALIDADES", campos.get("penalidades", ""))
    bloco_estruturado("CLÁUSULA DÉCIMA SEGUNDA – DA RESCISÃO", campos.get("rescisao", ""))
    bloco_estruturado("CLÁUSULA DÉCIMA TERCEIRA – DAS ALTERAÇÕES", campos.get("alteracoes", ""))
    bloco_estruturado("CLÁUSULA DÉCIMA QUARTA – DO FORO", campos.get("foro", ""))
    bloco_estruturado("CLÁUSULA DÉCIMA QUINTA – DISPOSIÇÕES GERAIS", campos.get("disposicoes_gerais", ""))
    
    # Assinaturas
    doc.add_page_break()
    doc.add_paragraph()
    doc.add_paragraph(f"São Paulo, {campos.get('data_assinatura', '__/__/____')}")
    doc.add_paragraph()
    doc.add_paragraph()
    
    assinatura_contratante = doc.add_paragraph("_" * 60)
    assinatura_contratante.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("CONTRATANTE").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Tribunal de Justiça do Estado de São Paulo").alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    assinatura_contratada = doc.add_paragraph("_" * 60)
    assinatura_contratada.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("CONTRATADA").alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Salvar buffer e disco
    nome_arquivo = f"Contrato_{campos.get('numero_contrato','TJSP-CONT').replace('/', '-')}.docx"
    exports_path = EXPORTS_DIR
    exports_path.mkdir(parents=True, exist_ok=True)
    caminho = str(exports_path / nome_arquivo)
    
    # Criar buffer primeiro
    buffer = io.BytesIO()
    try:
        doc.save(buffer)
        buffer.seek(0)
        print(f"[integration_contrato] Buffer DOCX criado: {len(buffer.getvalue())} bytes")
    except Exception as e:
        print(f"[integration_contrato] ERRO ao criar buffer: {e}")
        return None
    
    # Salvar no session_state se disponível
    if session_state is not None:
        try:
            session_state["contrato_docx_buffer"] = buffer
            session_state["contrato_docx_nome"] = nome_arquivo
            print(f"[integration_contrato] ✅ Buffer salvo no session_state")
        except Exception as e:
            print(f"[integration_contrato] ERRO session_state: {e}")
    
    # Tentar salvar no disco (opcional)
    try:
        doc.save(caminho)
        print(f"[integration_contrato] DOCX salvo em disco: {caminho}")
        return caminho
    except Exception:
        return None


def _gerar_contrato_docx_simples(campos: Dict[str, str], texto_completo: Optional[str] = None, session_state: dict = None) -> Optional[str]:
    """Versão simplificada sem formatação avançada (fallback)."""
    if Document is None:
        return None
    
    doc = Document()
    
    doc.add_heading('TRIBUNAL DE JUSTIÇA DO ESTADO DE SÃO PAULO', level=1)
    doc.add_heading(f"CONTRATO ADMINISTRATIVO Nº {campos.get('numero_contrato', 'XXXXX/YYYY')}", level=1)
    doc.add_paragraph(f"Data: {campos.get('data_assinatura', '__/__/____')}")
    doc.add_paragraph()
    
    def bloco(titulo: str, corpo: str):
        doc.add_heading(titulo, level=2)
        doc.add_paragraph(corpo if corpo else "(Não especificado)")
    
    bloco("CLÁUSULA PRIMEIRA – DO OBJETO", campos.get("objeto", ""))
    bloco("CLÁUSULA SEGUNDA – DA FUNDAMENTAÇÃO LEGAL", campos.get("fundamentacao_legal", ""))
    bloco("CLÁUSULA TERCEIRA – DA VIGÊNCIA", campos.get("vigencia", ""))
    bloco("CLÁUSULA QUARTA – DO VALOR GLOBAL", campos.get("valor_global", ""))
    bloco("CLÁUSULA QUINTA – DA FORMA DE PAGAMENTO", campos.get("forma_pagamento", ""))
    bloco("CLÁUSULA SEXTA – DO REAJUSTE", campos.get("reajuste", ""))
    bloco("CLÁUSULA SÉTIMA – DA GARANTIA", campos.get("garantia_contratual", ""))
    bloco("CLÁUSULA OITAVA – OBRIGAÇÕES DA CONTRATADA", campos.get("obrigacoes_contratada", ""))
    bloco("CLÁUSULA NONA – OBRIGAÇÕES DA CONTRATANTE", campos.get("obrigacoes_contratante", ""))
    bloco("CLÁUSULA DÉCIMA – FISCALIZAÇÃO", campos.get("fiscalizacao", ""))
    bloco("CLÁUSULA DÉCIMA PRIMEIRA – PENALIDADES", campos.get("penalidades", ""))
    bloco("CLÁUSULA DÉCIMA SEGUNDA – RESCISÃO", campos.get("rescisao", ""))
    bloco("CLÁUSULA DÉCIMA TERCEIRA – ALTERAÇÕES", campos.get("alteracoes", ""))
    bloco("CLÁUSULA DÉCIMA QUARTA – FORO", campos.get("foro", ""))
    bloco("CLÁUSULA DÉCIMA QUINTA – DISPOSIÇÕES GERAIS", campos.get("disposicoes_gerais", ""))
    
    # Salvar
    nome_arquivo = f"Contrato_{campos.get('numero_contrato','TJSP-CONT').replace('/', '-')}.docx"
    exports_path = EXPORTS_DIR
    exports_path.mkdir(parents=True, exist_ok=True)
    caminho = str(exports_path / nome_arquivo)
    
    buffer = io.BytesIO()
    try:
        doc.save(buffer)
        buffer.seek(0)
    except Exception as e:
        print(f"[integration_contrato] ERRO buffer simples: {e}")
        return None
    
    if session_state is not None:
        try:
            session_state["contrato_docx_buffer"] = buffer
            session_state["contrato_docx_nome"] = nome_arquivo
        except Exception:
            pass
    
    try:
        doc.save(caminho)
        return caminho
    except Exception:
        return None
