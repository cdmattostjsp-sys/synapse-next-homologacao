# -*- coding: utf-8 -*-
"""
utils/formatter_docx.py
======================================================
Conversão e formatação de relatórios institucionais TJSP/SAAB.
Responsável pela criação de artefatos DOCX oficiais dos módulos:
DFD, ETP, TR, Edital e Contrato.

Inclui:
- criação de documentos com cabeçalho institucional;
- assinatura técnica automatizada;
- compatibilidade retroativa com markdown_to_docx();
======================================================
"""

import os
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ===============================================================
# 🔧 Funções internas de formatação e cabeçalho
# ===============================================================

def criar_documento(titulo: str = "Relatório Institucional") -> Document:
    """
    Cria um documento DOCX com cabeçalho e metadados institucionais padrão SAAB/TJSP.
    """
    doc = Document()

    # Define margens e estilo institucional
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.18)
        section.right_margin = Inches(1.18)

    # Fonte padrão
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(11)

    # Cabeçalho institucional
    titulo_formatado = f"{titulo}\nSecretaria de Administração e Abastecimento – SAAB/TJSP"
    p = doc.add_paragraph(titulo_formatado)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True

    doc.add_paragraph("")  # espaço

    return doc


def adicionar_assinatura(doc: Document) -> None:
    """
    Adiciona assinatura técnica e data de geração institucional ao final do documento.
    """
    doc.add_paragraph("")
    data = datetime.now().strftime("%d/%m/%Y %H:%M")
    assinatura = (
        "\n\n____________________________________\n"
        "Gerado via SynapseNext vNext – SAAB/TJSP\n"
        f"{data}"
    )
    p = doc.add_paragraph(assinatura)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


# ===============================================================
# 🧩 Funções públicas de geração de relatórios
# ===============================================================

def gerar_relatorio_basico(conteudo: dict, output_path: str = "exports/relatorios/relatorio_institucional.docx") -> str:
    """
    Gera um relatório institucional básico a partir de um dicionário de dados.
    Cada chave vira um título e o valor vira o parágrafo correspondente.
    """
    doc = criar_documento("Relatório Técnico – SynapseNext vNext")

    for secao, texto in conteudo.items():
        doc.add_heading(str(secao), level=2)
        p = doc.add_paragraph(str(texto))
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        doc.add_paragraph("")

    adicionar_assinatura(doc)
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)

    print(f"✅ Relatório gerado com sucesso: {output_path}")
    return output_path


def gerar_relatorio_completo(conteudo: dict, titulo: str, output_path: str = "exports/relatorios/relatorio_completo.docx") -> str:
    """
    Gera um relatório com título customizado e múltiplas seções.
    """
    doc = criar_documento(titulo)

    for secao, texto in conteudo.items():
        doc.add_heading(secao, level=2)
        p = doc.add_paragraph(str(texto))
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        doc.add_paragraph("")

    adicionar_assinatura(doc)
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)

    print(f"✅ Relatório completo gerado: {output_path}")
    return output_path


# ===============================================================
# 🔁 Compatibilidade retroativa – suporte a Markdown
# ===============================================================

def markdown_to_docx(markdown_text: str, output_path: str = "exports/relatorios/markdown_export.docx") -> str:
    """
    Converte um texto Markdown em DOCX, preservando o estilo institucional SAAB/TJSP.
    """
    doc = criar_documento("Conversão Markdown – SynapseNext vNext")

    lines = markdown_text.split("\n")
    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Títulos Markdown
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        # Listas
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif re.match(r"^\d+\.\s", line):
            doc.add_paragraph(re.sub(r"^\d+\.\s", "", line).strip(), style="List Number")
        # Parágrafos
        else:
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    adicionar_assinatura(doc)
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"✅ Documento Markdown convertido: {output_path}")
    return output_path


# ===============================================================
# 🔍 Execução direta (teste local)
# ===============================================================

if __name__ == "__main__":
    exemplo = {
        "Contexto": "Este é um relatório gerado para teste da formatação institucional SAAB/TJSP.",
        "Objetivo": "Validar a consistência de fontes, margens e cabeçalhos.",
        "Resultado": "O sistema SynapseNext vNext está funcional e padronizado."
    }
    gerar_relatorio_basico(exemplo, "exports/relatorios/teste_formatter_docx.docx")
