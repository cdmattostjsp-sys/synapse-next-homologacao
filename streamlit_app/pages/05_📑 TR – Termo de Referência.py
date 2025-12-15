import sys
from pathlib import Path
ROOT = Path(__file__).resolve().parents[2]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

# ==============================
# pages/05_📑 TR – Termo de Referência.py
# SynapseNext – Secretaria de Administração e Abastecimento (TJSP)
# ==============================

import streamlit as st
from datetime import datetime
import os, sys, json
from io import BytesIO
from docx import Document
from utils.ui_components import aplicar_estilo_global, exibir_cabecalho_padrao
from home_utils.sidebar_organizer import apply_sidebar_grouping
from utils.integration_tr import export_tr_to_json, ler_modelos_tr
from home_utils.refinamento_ia import render_refinamento_iterativo

# ==========================================================
# 🔄 Lazy Loading da OpenAI Client
# ==========================================================
def _get_openai_client():
    """Carrega OpenAI client sob demanda (lazy loading)."""
    try:
        from openai import OpenAI
        api_key = os.getenv("OPENAI_API_KEY") or os.getenv("openai_api_key")
        if not api_key:
            return None
        return OpenAI(api_key=api_key)
    except Exception as e:
        st.error(f"⚠️ Erro ao carregar OpenAI client: {e}")
        return None

# ==========================================================
# ⚙️ Configuração
# ==========================================================
st.set_page_config(page_title="📑 Termo de Referência", layout="wide", page_icon="📑")
apply_sidebar_grouping()

# Estilo institucional PJe-inspired
st.markdown("""
<style>
/* ============================================
   PADRÃO VISUAL PJe-INSPIRED - SYNAPSE NEXT
   Versão: 2025.1-homolog
   ============================================ */

/* Título principal - tamanho reduzido para sobriedade */
h1 {
    font-size: 1.8rem !important;
    font-weight: 500 !important;
    color: #2c3e50 !important;
    margin-bottom: 0.3rem !important;
}

/* Caption institucional */
.caption {
    color: #6c757d;
    font-size: 0.9rem;
    margin-bottom: 1rem;
}

/* Bloco de IA - destaque sutil */
.ia-block {
    border: 1px solid #d0d7de;
    border-radius: 3px;
    padding: 1rem 1.2rem;
    background-color: #f0f2f5;
    margin: 1rem 0 1.2rem 0;
}
.ia-block h3 {
    font-size: 1rem;
    font-weight: 600;
    color: #1f2937;
    margin: 0 0 0.6rem 0;
    letter-spacing: -0.01em;
}

/* Seções com fundo cinza - contraste melhorado */
h3 {
    font-size: 1.1rem !important;
    font-weight: 500 !important;
    color: #374151 !important;
    background-color: #e5e7eb !important;
    padding: 0.6rem 0.8rem !important;
    border-radius: 3px !important;
    margin-top: 1.5rem !important;
    margin-bottom: 1rem !important;
}

/* Botões - destaque apenas para ações principais */
div.stButton > button {
    border-radius: 3px;
    font-weight: 500;
    border: 1px solid #d0d7de;
}
div.stButton > button[kind="primary"] {
    background-color: #0969da !important;
    border-color: #0969da !important;
}

/* Formulário clean */
.stTextInput label, .stTextArea label {
    font-weight: 500;
    color: #1f2937;
    font-size: 0.9rem;
}

/* Expander refinamento com destaque discreto */
details {
    border: 1px solid #d0d7de;
    border-radius: 3px;
    padding: 0.5rem;
    background-color: #ffffff;
}
summary {
    font-weight: 500;
    color: #0969da;
    cursor: pointer;
}
</style>
""", unsafe_allow_html=True)

# ==========================================================
# 🏛️ Cabeçalho institucional
# ==========================================================
st.markdown("<h1>Termo de Referência (TR)</h1>", unsafe_allow_html=True)
st.markdown("<p class='caption'>Preencha manualmente ou processe insumos com IA especializada</p>", unsafe_allow_html=True)
st.divider()

# ==========================================================
# 🔍 Carregamento inteligente: INSUMOS → TRAgent → Formulário
# ==========================================================

# Paths dos arquivos
INSUMO_TR_PATH = os.path.join("exports", "insumos", "json", "TR_ultimo.json")
TR_JSON_PATH = os.path.join("exports", "tr_data.json")

# ==========================================================
# 📝 Definição das seções do TR para refinamento iterativo
# ==========================================================
SECOES_TR = [
    "objeto",
    "justificativa_tecnica",
    "especificacao_tecnica",
    "criterios_julgamento",
    "riscos",
    "observacoes_finais",
    "prazo_execucao",
    "estimativa_valor",
    "fonte_recurso"
]

tr_salvo = {}
insumo_detectado = False

# 1️⃣ Verificar se existe TR processado
if os.path.exists(TR_JSON_PATH):
    try:
        with open(TR_JSON_PATH, "r", encoding="utf-8") as f:
            dados_tr = json.load(f)
            tr_salvo = dados_tr.get("TR", {})
            if tr_salvo and any(v for v in tr_salvo.values() if v):
                st.success("📎 TR processado carregado (exports/tr_data.json)")
    except Exception as e:
        st.warning(f"⚠️ Erro ao carregar TR: {e}")

# 2️⃣ Se não houver TR processado, verificar INSUMO bruto
if not tr_salvo or not any(v for v in tr_salvo.values() if v):
    if os.path.exists(INSUMO_TR_PATH):
        insumo_detectado = True
        st.info("📄 Insumo TR detectado. Use o botão **'Processar com IA'** para extrair as 9 seções automaticamente.")
        
        # Carregar texto bruto do insumo (para preview)
        try:
            with open(INSUMO_TR_PATH, "r", encoding="utf-8") as f:
                insumo_data = json.load(f)
                texto_bruto = insumo_data.get("conteudo_textual", "")
                
                if texto_bruto and len(texto_bruto) > 100:
                    with st.expander("👁️ Preview do insumo carregado", expanded=False):
                        st.text_area(
                            "Texto extraído do PDF/DOCX:",
                            texto_bruto[:1000] + "..." if len(texto_bruto) > 1000 else texto_bruto,
                            height=200,
                            disabled=True
                        )
        except Exception as e:
            st.warning(f"⚠️ Erro ao ler insumo: {e}")

# 3️⃣ Nenhum dado disponível
if not tr_salvo and not insumo_detectado:
    st.info("ℹ️ Nenhum TR ou insumo detectado. Faça upload no módulo **🔧 Insumos** ou preencha manualmente.")

# ==========================================================
# 🧾 Formulário TR – 9 Seções Estruturadas
# ==========================================================
st.markdown("### Assistente IA")
st.caption("Processamento automático: requer insumos do módulo anterior")

col_ia1, col_ia2, col_ia3 = st.columns(3)

with col_ia1:
    if st.button("⚡ Processar com IA", use_container_width=True, type="primary", key="btn_ia_processar"):
        try:
            with st.spinner("Processando TR com IA especializada..."):
                from utils.integration_tr import gerar_tr_com_ia
                resultado = gerar_tr_com_ia()
                
            if "erro" in resultado:
                st.error(f"❌ {resultado['erro']}")
            else:
                st.success("TR estruturado com sucesso")
                
                # Exibir resumo
                tr_processado = resultado.get("TR", {})
                secoes_ia = sum(1 for v in tr_processado.values() if v and v.strip())
                
                col_a, col_b, col_c = st.columns(3)
                with col_a:
                    st.metric("Seções", f"{secoes_ia}/9")
                with col_b:
                    st.metric("Prazo", tr_processado.get("prazo_execucao", "N/A")[:20])
                with col_c:
                    st.metric("Valor", f"R$ {tr_processado.get('estimativa_valor', '0,00')}")
                
                st.rerun()
        except Exception as e:
            st.error(f"Erro ao processar: {e}")

with col_ia2:
    if st.button("📤 Enviar para Edital", use_container_width=True, disabled=not tr_salvo, key="btn_enviar_edital"):
        try:
            from datetime import datetime
            
            base = os.path.join("exports", "insumos", "json")
            os.makedirs(base, exist_ok=True)
            
            payload = {
                "artefato": "EDITAL",
                "origem": "TR_estruturado",
                "data_processamento": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "status": "ok",
                "campos_ai": tr_salvo,
                "conteudo_textual": "",
            }
            
            arq_ultimo = os.path.join(base, "EDITAL_ultimo.json")
            with open(arq_ultimo, "w", encoding="utf-8") as f:
                json.dump(payload, f, ensure_ascii=False, indent=2)
            
            st.success("Dados enviados para o módulo Edital")
            st.info("Acesse o módulo Edital para continuar")
            
        except Exception as e:
            st.error(f"Erro: {e}")

with col_ia3:
    st.write("")  # Espaçamento

st.divider()

# Refinamento iterativo
tr_salvo = render_refinamento_iterativo(
    secoes_disponiveis=SECOES_TR,
    dados_atuais=tr_salvo if tr_salvo else {},
    artefato="TR",
    campos_simples=None
)

st.divider()

# ==========================================================
# Formulário com 9 seções estruturadas
# ==========================================================
st.markdown("### Seções do Termo de Referência")

# Seção 1: Objeto
objeto = st.text_area(
    "1. Objeto da Contratação",
    value=tr_salvo.get("objeto", ""),
    height=120,
    key="tr_objeto",
    help="Descrição do objeto a ser contratado"
)

# Seção 2: Justificativa Técnica (coluna única)
justificativa_tecnica = st.text_area(
    "2. Justificativa Técnica",
    value=tr_salvo.get("justificativa_tecnica", ""),
    height=150,
    key="tr_just",
    help="Fundamentação da necessidade da contratação"
)

# Seção 3: Especificações Técnicas (coluna única)
especificacao_tecnica = st.text_area(
    "3. Especificações Técnicas",
    value=tr_salvo.get("especificacao_tecnica", ""),
    height=150,
    key="tr_espec",
    help="Detalhamento técnico dos serviços/produtos"
)

# Seção 4: Critérios de Julgamento (coluna única)
criterios_julgamento = st.text_area(
    "4. Critérios de Julgamento",
    value=tr_salvo.get("criterios_julgamento", ""),
    height=120,
    key="tr_crit",
    help="Critérios para avaliação das propostas"
)

# Seção 5: Riscos (coluna única)
riscos = st.text_area(
    "5. Riscos Associados",
    value=tr_salvo.get("riscos", ""),
    height=120,
    key="tr_riscos",
    help="Identificação e mitigação de riscos"
)

# Seção 6: Observações Finais (coluna única)
observacoes_finais = st.text_area(
    "6. Observações Finais",
    value=tr_salvo.get("observacoes_finais", ""),
    height=120,
    key="tr_obs",
    help="Informações complementares e observações"
)

st.divider()

# ==========================================================
# Campos complementares (Seções 7-9)
# ==========================================================
st.markdown("### 📊 Informações Complementares")

col3, col4, col5 = st.columns(3)
with col3:
    prazo_execucao = st.text_input(
        "7. Prazo de Execução",
        value=tr_salvo.get("prazo_execucao", ""),
        key="tr_prazo",
        help="Prazo estimado para execução"
    )
with col4:
    estimativa_valor = st.text_input(
        "8. Estimativa de Valor (R$)",
        value=tr_salvo.get("estimativa_valor", ""),
        key="tr_valor",
        help="Valor estimado da contratação"
    )
with col5:
    fonte_recurso = st.text_input(
        "9. Fonte de Recurso",
        value=tr_salvo.get("fonte_recurso", ""),
        key="tr_fonte",
        help="Origem do recurso orçamentário"
    )

st.divider()

# Botões de ação
col_salvar, col_baixar = st.columns(2)

with col_salvar:
    if st.button("Salvar TR", type="secondary", use_container_width=True):
        tr_completo = {
            "objeto": objeto,
            "justificativa_tecnica": justificativa_tecnica,
            "especificacao_tecnica": especificacao_tecnica,
            "criterios_julgamento": criterios_julgamento,
            "riscos": riscos,
            "observacoes_finais": observacoes_finais,
            "prazo_execucao": prazo_execucao,
            "estimativa_valor": estimativa_valor,
            "fonte_recurso": fonte_recurso,
        }
        
        resultado = export_tr_to_json(tr_completo)
        if "erro" not in resultado:
            st.success("TR salvo com sucesso")
        else:
            st.error(f"Erro: {resultado['erro']}")

with col_baixar:
    if st.button("Baixar TR (DOCX)", use_container_width=True):
        doc = Document()
        
        doc.add_heading("Termo de Referência (TR)", level=1)
        doc.add_paragraph("Lei 14.133/2021 - Nova Lei de Licitações")
        
        # Seções do TR
        doc.add_heading("Seções do Termo de Referência", level=2)
        
        secoes_nomes = [
            ("1. Objeto da Contratação", objeto),
            ("2. Justificativa Técnica", justificativa_tecnica),
            ("3. Especificações Técnicas", especificacao_tecnica),
            ("4. Critérios de Julgamento", criterios_julgamento),
            ("5. Riscos Associados", riscos),
            ("6. Observações Finais", observacoes_finais),
        ]
        
        for nome, conteudo in secoes_nomes:
            doc.add_heading(nome, level=3)
            if conteudo and conteudo.strip():
                doc.add_paragraph(conteudo)
            else:
                doc.add_paragraph("[Não preenchido]")
        
        # Informações Complementares
        doc.add_heading("Informações Complementares", level=2)
        doc.add_paragraph(f"7. Prazo de Execução: {prazo_execucao if prazo_execucao else '[Não preenchido]'}")
        doc.add_paragraph(f"8. Estimativa de Valor: R$ {estimativa_valor if estimativa_valor else '[Não preenchido]'}")
        doc.add_paragraph(f"9. Fonte de Recurso: {fonte_recurso if fonte_recurso else '[Não preenchido]'}")
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        st.download_button(
            label="⬇️ Download DOCX (TR completo)",
            data=buffer,
            file_name="TR_completo_Lei14133.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )

st.caption("Dica: Use o Assistente IA para preencher automaticamente as seções a partir dos insumos carregados")
