# ==========================================================
# pages/07_📃 Contrato.py – SynapseNext / SAAB TJSP
# ==========================================================
# Módulo final da jornada de contratação pública.
# Gera a minuta do contrato a partir de insumos cumulativos
# (DFD, ETP, TR, Edital) e processamento IA institucional.
# ==========================================================

import os, json
from io import BytesIO
from datetime import datetime
import streamlit as st
from docx import Document
from openai import OpenAI
from pathlib import Path
from utils.ui_components import aplicar_estilo_global, exibir_cabecalho_padrao

# ==========================================================
# ⚙️ Configuração básica
# ==========================================================
st.set_page_config(page_title="📃 Contrato", layout="wide", page_icon="📃")
aplicar_estilo_global()

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY") or os.getenv("openai_api_key")
client = OpenAI(api_key=OPENAI_API_KEY)

# ==========================================================
# 📚 Leitura de modelos institucionais
# ==========================================================
def ler_modelos_contrato() -> str:
    base = Path(__file__).resolve().parents[1] / "knowledge" / "contrato_models"
    textos = []
    if base.exists():
        for arq in base.glob("*.txt"):
            try:
                textos.append(arq.read_text(encoding="utf-8"))
            except Exception:
                pass
    return "\n\n".join(textos)

# ==========================================================
# 🏛️ Cabeçalho institucional
# ==========================================================
exibir_cabecalho_padrao(
    "📃 Minuta do Contrato Administrativo",
    "Consolidação final dos artefatos da jornada de contratação (DFD → ETP → TR → Edital → Contrato)"
)
st.divider()

# ==========================================================
# 🔗 Dados cumulativos disponíveis
# ==========================================================
defaults = {}
for chave in ["dfd_campos_ai", "etp_campos_ai", "tr_campos_ai", "edital_campos_ai", "contrato_campos_ai"]:
    if chave in st.session_state:
        defaults.update(st.session_state[chave])

if defaults:
    st.success("📎 Dados recebidos automaticamente dos módulos anteriores (DFD, ETP, TR, Edital).")
else:
    st.info("Nenhum insumo ativo detectado. Você pode preencher manualmente ou aguardar integração via módulo INSUMOS.")

# ==========================================================
# 🧾 Formulário – Campos contratuais
# ==========================================================
st.subheader("📄 Dados do Contrato")

col1, col2 = st.columns(2)
with col1:
    objeto = st.text_area("Objeto do Contrato", value=defaults.get("objeto", ""), height=100)
    partes = st.text_area("Partes Contratantes", value=defaults.get("partes", ""), height=80)
    vigencia = st.text_input("Vigência", value=defaults.get("vigencia", "12 meses a contar da assinatura"))
    valor_global = st.text_input("Valor Global", value=defaults.get("valor_global", ""))
    reajuste = st.text_area("Reajuste", value=defaults.get("reajuste", "Conforme índice oficial e cláusulas legais"), height=70)
    garantias = st.text_area("Garantias", value=defaults.get("garantias", ""), height=70)

with col2:
    prazos_pagamento = st.text_area("Prazos e Forma de Pagamento", value=defaults.get("prazos_pagamento", ""), height=70)
    obrigacoes_contratada = st.text_area("Obrigações da Contratada", value=defaults.get("obrigacoes_contratada", ""), height=100)
    obrigacoes_contratante = st.text_area("Obrigações da Contratante", value=defaults.get("obrigacoes_contratante", ""), height=100)
    fiscalizacao = st.text_area("Fiscalização e Acompanhamento", value=defaults.get("fiscalizacao", ""), height=70)
    penalidades = st.text_area("Penalidades", value=defaults.get("penalidades", ""), height=80)
    rescisao = st.text_area("Rescisão Contratual", value=defaults.get("rescisao", ""), height=80)
    foro = st.text_input("Foro Competente", value=defaults.get("foro", "Comarca de São Paulo/SP"))

st.divider()
observacoes_finais = st.text_area("Observações Finais", value=defaults.get("observacoes_finais", ""), height=70)

# ==========================================================
# ⚙️ Geração com IA Institucional
# ==========================================================
st.subheader("⚙️ Geração da Minuta Contratual com IA Institucional")

if st.button("🤖 Gerar minuta completa do Contrato com IA institucional"):
    with st.spinner("Gerando minuta contratual com base nos artefatos e modelos institucionais..."):
        modelos = ler_modelos_contrato()
        campos = {
            "objeto": objeto,
            "partes": partes,
            "vigencia": vigencia,
            "valor_global": valor_global,
            "reajuste": reajuste,
            "garantias": garantias,
            "prazos_pagamento": prazos_pagamento,
            "obrigacoes_contratada": obrigacoes_contratada,
            "obrigacoes_contratante": obrigacoes_contratante,
            "fiscalizacao": fiscalizacao,
            "penalidades": penalidades,
            "rescisao": rescisao,
            "foro": foro,
            "observacoes_finais": observacoes_finais
        }

        user_prompt = f"""
Com base nos campos abaixo e nos modelos institucionais da SAAB/TJSP,
elabore a minuta completa de um CONTRATO ADMINISTRATIVO conforme a Lei nº 14.133/2021.
O texto deve seguir o padrão redacional do TJSP.

Campos:
{json.dumps(campos, ensure_ascii=False, indent=2)}

Modelos de referência:
\"\"\"{modelos}\"\"\"
"""

        try:
            response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": "Você é um redator institucional do TJSP responsável por elaborar contratos administrativos conforme o padrão SAAB/TJSP e a Lei 14.133/2021."},
                    {"role": "user", "content": user_prompt}
                ],
                temperature=0.3
            )
            artefato_contrato = response.choices[0].message.content.strip()
            st.session_state["artefato_contrato_gerado"] = artefato_contrato
            st.success("✅ Minuta contratual gerada com sucesso! Você pode visualizar e exportar o documento abaixo.")
            st.text_area("📄 Pré-visualização da minuta gerada:", artefato_contrato, height=400)

        except Exception as e:
            st.error(f"Erro ao gerar minuta contratual com IA: {e}")

# ==========================================================
# 💾 Exportação DOCX
# ==========================================================
if "artefato_contrato_gerado" in st.session_state:
    doc = Document()
    doc.add_heading("MINUTA DO CONTRATO ADMINISTRATIVO", level=1)
    doc.add_paragraph(st.session_state["artefato_contrato_gerado"])

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    st.download_button(
        label="📤 Exportar contrato em DOCX",
        data=buffer,
        file_name=f"Contrato_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

st.caption("📎 O texto acima é gerado pela IA institucional com base nos modelos oficiais do TJSP e nos artefatos acumulados da jornada de contratação.")
