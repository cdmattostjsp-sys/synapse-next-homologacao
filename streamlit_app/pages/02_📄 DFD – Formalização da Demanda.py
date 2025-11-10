# ==========================================================
# pages/02_📄 DFD – Formalização da Demanda.py
# SynapseNext – Secretaria de Administração e Abastecimento (TJSP)
# Revisão: Engenheiro Synapse – vNext_2025.11.09-r3
# Correção consolidada: leitura, mapeamento e pré-preenchimento DFD
# ==========================================================

import json
import re
from pathlib import Path
from io import BytesIO
import streamlit as st
from docx import Document

# ==========================================================
# 📦 Imports institucionais
# ==========================================================
from utils.agents_bridge import AgentsBridge
from utils.ui_components import aplicar_estilo_global, exibir_cabecalho_padrao

# ==========================================================
# ⚙️ Configuração inicial
# ==========================================================
st.set_page_config(
    page_title="📄 DFD – Formalização da Demanda",
    layout="wide",
    page_icon="📄",
)
aplicar_estilo_global()

exibir_cabecalho_padrao(
    "📄 Formalização da Demanda (DFD)",
    "Registro institucional da demanda e geração de rascunho com IA",
)
st.divider()

# ==========================================================
# 🔹 Mapeamento semântico do JSON da IA
# ==========================================================
def mapear_dfd_campos(dados_ia: dict) -> dict:
    """Transforma o JSON complexo retornado pela IA em um dicionário plano."""
    campos = {}
    dfd = dados_ia.get("DFD", {}) or dados_ia.get("estudo_tecnico_preliminar", {})

    campos["lei"] = dfd.get("lei", "")
    campos["processo_cpa"] = dfd.get("processo_cpa", "")

    objetivo = dfd.get("objetivo", {})
    campos["descricao_necessidade"] = objetivo.get("contratacao", "")
    campos["localizacao"] = objetivo.get("localizacao", {}).get("descricao", "")
    campos["endereco"] = objetivo.get("localizacao", {}).get("endereco", "")
    campos["disciplinas"] = ", ".join(objetivo.get("disciplinas", []))

    necessidade = dfd.get("descricao_da_necessidade", {})
    edificio = necessidade.get("edificio", {})
    campos["caracteristicas_edificio"] = f"{edificio.get('pavimentos', '')} pavimentos, {edificio.get('sistema_construtivo', '')}"
    campos["intervencoes_previstas"] = ", ".join(necessidade.get("intervencoes", []))

    plano = dfd.get("previsto_no_plano_de_contratacoes_anual", {}).get("plano_obras", {})
    campos["ano_plano_obras"] = plano.get("ano", "")
    campos["codigo_pca"] = dfd.get("previsto_no_plano_de_contratacoes_anual", {}).get("codigo_identificacao", "")

    planejamento = dfd.get("planejamento_estrategico", {})
    campos["periodo_planejamento"] = planejamento.get("periodo", "")
    campos["objetivos_estrategicos"] = ", ".join(planejamento.get("objetivos", []))

    return campos


# ==========================================================
# 🧩 Funções robustas de carregamento
# ==========================================================
def _carregar_insumo_dfd() -> dict:
    """Carrega o último insumo processado e converte o JSON gerado pela IA."""
    candidatos = [
        Path("exports") / "insumos" / "json" / "DFD_ultimo.json",
        Path("/workspaces/synapse-next-homologacao/exports/insumos/json/DFD_ultimo.json"),
    ]
    base_path = next((p for p in candidatos if p.exists()), None)
    if not base_path:
        st.info("ℹ️ Nenhum insumo DFD encontrado. Gere um pelo módulo 'Insumos'.")
        return {}

    try:
        with open(base_path, "r", encoding="utf-8") as f:
            dados = json.load(f)

        if "resultado_ia" in dados:
            resposta = dados["resultado_ia"].get("resposta_texto", "")
            conteudo_json = None

            # 🧩 Extrai o bloco JSON
            match = re.search(r"```json(.*?)```", resposta, re.S)
            if match:
                conteudo_json = match.group(1).strip()
            elif "{" in resposta and "}" in resposta:
                start = resposta.find("{")
                end = resposta.rfind("}") + 1
                conteudo_json = resposta[start:end].strip()

            if conteudo_json:
                try:
                    dados_ia = json.loads(conteudo_json)

                    # 🪞 Se houver encapsulamento "DFD", mergulha nele
                    if isinstance(dados_ia, dict) and "DFD" in dados_ia:
                        dados_ia = dados_ia["DFD"]

                    return mapear_dfd_campos(dados_ia)

                except json.JSONDecodeError:
                    st.warning("⚠️ JSON parcial detectado, tentando normalizar...")
                    conteudo_json = conteudo_json.replace("\n", " ").replace("```json", "").replace("```", "")
                    open_braces = conteudo_json.count("{")
                    close_braces = conteudo_json.count("}")
                    if open_braces > close_braces:
                        conteudo_json += "}" * (open_braces - close_braces)
                    conteudo_json = conteudo_json[:conteudo_json.rfind("}") + 1]
                    dados_ia = json.loads(conteudo_json)
                    st.success("✅ JSON parcial recuperado.")
                    return mapear_dfd_campos(dados_ia)

        return dados.get("campos_ai", {})

    except Exception as e:
        st.error(f"❌ Erro ao carregar insumo DFD ({e})")
        return {}


def _carregar_lacunas_dfd() -> list[str]:
    """Carrega lacunas (itens não inferidos pela IA)."""
    candidatos = [
        Path("exports") / "insumos" / "json" / "DFD_ultimo.json",
        Path("/workspaces/synapse-next-homologacao/exports/insumos/json/DFD_ultimo.json"),
    ]
    base_path = next((p for p in candidatos if p.exists()), None)
    if not base_path:
        return []
    try:
        with open(base_path, "r", encoding="utf-8") as f:
            dados = json.load(f)
        return dados.get("lacunas", []) or dados.get("campos_ai", {}).get("lacunas", [])
    except Exception:
        return []


# ==========================================================
# 🔍 Carregar dados inferidos pela IA (ordem correta)
# ==========================================================
campos_ai = _carregar_insumo_dfd()
lacunas_ai = _carregar_lacunas_dfd()

# ==========================================================
# 🧾 Formulário DFD – pré-preenchido
# ==========================================================
st.subheader("1️⃣ Entrada – Formalização da Demanda")

descricao_default = campos_ai.get("descricao_necessidade", "")
motivacao_default = campos_ai.get("objetivos_estrategicos", "")
prazo_default = campos_ai.get("periodo_planejamento", "")
estimativa_default = 0.0

with st.form("form_dfd"):
    col1, col2 = st.columns(2)

    with col1:
        unidade = st.text_input("Unidade Demandante", placeholder="Ex.: Secretaria de Administração e Abastecimento – SAAB", key="dfd_unidade")
        responsavel = st.text_input("Responsável pela Demanda", key="dfd_responsavel")
        prazo = st.text_input("Prazo Estimado para Atendimento", value=prazo_default, key="dfd_prazo")

    with col2:
        descricao = st.text_area("Descrição da Necessidade", height=100, value=descricao_default, key="dfd_descricao")
        motivacao = st.text_area("Motivação / Objetivos Estratégicos", height=100, value=motivacao_default, key="dfd_motivacao")
        estimativa_valor = st.number_input("Estimativa de Valor (R$)", min_value=0.0, step=1000.0, value=estimativa_default, key="dfd_estimativa_valor")

    colb1, colb2 = st.columns(2)
    with colb1:
        gerar_ia = st.form_submit_button("⚙️ Gerar rascunho com IA institucional")
    with colb2:
        salvar_manual = st.form_submit_button("💾 Salvar dados manualmente")

# ==========================================================
# 📝 Mostrar lacunas
# ==========================================================
if lacunas_ai:
    with st.expander("⚠️ Campos não inferidos pela IA"):
        for item in lacunas_ai:
            st.markdown(f"- {item}")

# ==========================================================
# 🤖 Execução da IA institucional
# ==========================================================
if gerar_ia:
    st.info("Executando agente DFD institucional com base no insumo processado...")
    metadata = {
        "unidade": unidade,
        "responsavel": responsavel,
        "prazo": prazo,
        "descricao": descricao,
        "motivacao": motivacao,
        "estimativa_valor": estimativa_valor,
        "campos_ai": campos_ai,
        "origem": "pagina_dfd_streamlit",
    }
    try:
        bridge = AgentsBridge("DFD")
        resultado = bridge.generate(metadata)
        st.success("✅ Rascunho gerado com sucesso.")
        st.json(resultado)
        st.session_state["last_dfd"] = resultado
    except Exception as e:
        st.error(f"Erro ao gerar rascunho com IA: {e}")

# ==========================================================
# 💾 Exportação
# ==========================================================
if "last_dfd" in st.session_state and st.session_state["last_dfd"]:
    st.divider()
    st.subheader("📤 Exportação de Documento")

    dfd_data = st.session_state["last_dfd"]
    doc = Document()
    doc.add_heading("Formalização da Demanda (DFD)", level=1)
    for k, v in dfd_data.items():
        p = doc.add_paragraph()
        p.add_run(f"{k}: ").bold = True
        p.add_run(str(v) or "—")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    st.download_button("💾 Baixar DFD_rascunho.docx", buffer, file_name="DFD_rascunho.docx")

st.caption("💡 Este módulo aceita preenchimento manual, mas dá prioridade ao insumo pré-processado pelo módulo INSUMOS + IA institucional.")
