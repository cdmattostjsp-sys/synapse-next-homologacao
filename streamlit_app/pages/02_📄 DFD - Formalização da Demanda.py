from __future__ import annotations

import sys
from pathlib import Path
ROOT = Path(__file__).resolve().parents[2]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

# -*- coding: utf-8 -*-
# ======================================================================
# pages/02_📄 DFD.py — VERSÃO FINAL 2025-D8 (ESTÁVEL)
# Formalização da Demanda (DFD) – Modelo Moderno-Governança
# Compatível com:
#   - utils/integration_dfd.py (2025-D8)
#   - agents/document_agent.py (D2)
#   - utils/ai_client.py vNext
# Fluxo: INSUMOS → DFD (formulário) → IA → DOCX
# ======================================================================

import json
import io
from typing import Any, Dict, List
from datetime import datetime

import streamlit as st
from home_utils.sidebar_organizer import apply_sidebar_grouping
from home_utils.refinamento_ia import render_refinamento_iterativo
from docx import Document

from utils.integration_dfd import (
    obter_dfd_da_sessao,
    salvar_dfd_em_json,
    gerar_rascunho_dfd_com_ia,
    status_dfd,
)

# ======================================================================
# ⚙️ CONFIGURAÇÃO DA PÁGINA
# ======================================================================
st.set_page_config(
    page_title="📄 Formalização da Demanda (DFD)",
    layout="wide",
)
apply_sidebar_grouping()

# Estilo institucional PJe-inspired
st.markdown("""
<style>
/* Título principal - tamanho reduzido para sobriedade */
h1 {
    font-size: 1.8rem !important;
    font-weight: 500 !important;
    color: #2c3e50 !important;
    margin-bottom: 0.3rem !important;
}
/* Caption institucional */
.caption {
    color: #6c757d;
    font-size: 0.9rem;
    margin-bottom: 1rem;
}
/* Bloco de IA - destaque sutil */
.ia-block {
    border: 1px solid #d0d7de;
    border-radius: 3px;
    padding: 1rem 1.2rem;
    background-color: #f6f8fa;
    margin: 1rem 0 1.2rem 0;
}
.ia-block h3 {
    font-size: 1rem;
    font-weight: 600;
    color: #1f2937;
    margin: 0 0 0.6rem 0;
    letter-spacing: -0.01em;
}
/* Seções com fundo cinza leve */
h3 {
    font-size: 1.1rem !important;
    font-weight: 500 !important;
    color: #374151 !important;
    background-color: #f3f4f6 !important;
    padding: 0.6rem 0.8rem !important;
    border-radius: 3px !important;
    margin-top: 1.5rem !important;
    margin-bottom: 1rem !important;
}
/* Botões - destaque apenas para ações principais */
div.stButton > button {
    border-radius: 3px;
    font-weight: 500;
    border: 1px solid #d0d7de;
}
div.stButton > button[kind="primary"] {
    background-color: #0969da !important;
    border-color: #0969da !important;
}
/* Formulário clean */
.stTextInput label, .stTextArea label {
    font-weight: 500;
    color: #1f2937;
    font-size: 0.9rem;
}
/* Expander refinamento com destaque discreto */
details {
    border: 1px solid #d0d7de;
    border-radius: 3px;
    padding: 0.5rem;
    background-color: #ffffff;
}
summary {
    font-weight: 500;
    color: #0969da;
    cursor: pointer;
}
</style>
""", unsafe_allow_html=True)

st.markdown("<h1>Formalização da Demanda (DFD)</h1>", unsafe_allow_html=True)
st.markdown("<p class='caption'>Preencha manualmente ou carregue dados do módulo Insumos</p>", unsafe_allow_html=True)
st.info(status_dfd())

# ======================================================================
# 📚 Constantes – padrão Moderno-Governança (11 seções)
# ======================================================================
SECOES_DFD: List[str] = [
    "Contexto Institucional",
    "Diagnóstico da Situação Atual",
    "Fundamentação da Necessidade",
    "Objetivos da Contratação",
    "Escopo Inicial da Demanda",
    "Resultados Esperados",
    "Benefícios Institucionais",
    "Justificativa Legal",
    "Riscos da Não Contratação",
    "Requisitos Mínimos",
    "Critérios de Sucesso",
]

# ======================================================================
# 🔧 Funções utilitárias
# ======================================================================
def _to_str(value: Any) -> str:
    """Converte qualquer estrutura em string legível para edição."""
    if value is None:
        return ""
    if isinstance(value, (dict, list)):
        try:
            return json.dumps(value, ensure_ascii=False, indent=2)
        except Exception:
            return str(value)
    return str(value)


def _extrair_secoes(dados_brutos: Dict[str, Any]) -> Dict[str, str]:
    """
    Extrai as 11 seções padrão do DFD Moderno-Governança.
    Se não existirem, retorna dicionário com chaves padrão vazias.
    """
    secoes_orig = dados_brutos.get("secoes")
    secoes_final: Dict[str, str] = {}

    if not isinstance(secoes_orig, dict):
        secoes_orig = {}

    for nome in SECOES_DFD:
        valor = secoes_orig.get(nome, "")
        if not isinstance(valor, str):
            valor = _to_str(valor)
        secoes_final[nome] = valor.strip()

    return secoes_final


def _extrair_lacunas(dados_brutos: Dict[str, Any]) -> List[str]:
    lac = dados_brutos.get("lacunas", [])
    if isinstance(lac, list):
        return [str(x) for x in lac]
    return []


def _extrair_admin_e_campos_tradicionais(dados_brutos: Dict[str, Any]) -> Dict[str, str]:
    """
    Extrai os campos administrativos e os dois campos tradicionais
    (descrição + motivação), mantendo compatibilidade com formatos antigos.
    """
    campos = dados_brutos or {}
    if not isinstance(campos, dict):
        campos = {}

    secoes = campos.get("secoes") if isinstance(campos.get("secoes"), dict) else {}

    # ------------------------------------------------------------
    # CAMPOS ADMINISTRATIVOS
    # ------------------------------------------------------------
    unidade = (
        campos.get("unidade_demandante")
        or campos.get("unidade")
        or ""
    )
    responsavel = campos.get("responsavel", "")
    prazo = (
        campos.get("prazo_estimado")
        or campos.get("prazo")
        or ""
    )
    valor_estimado = (
        campos.get("valor_estimado")
        or campos.get("estimativa_valor")
        or "0,00"
    )

    # ------------------------------------------------------------
    # DESCRIÇÃO DA NECESSIDADE
    # (Contexto + Diagnóstico + Fundamentação)
    # ------------------------------------------------------------
    descricao_txt = ""

    if isinstance(campos.get("descricao_necessidade"), str) and campos["descricao_necessidade"].strip():
        descricao_txt = campos["descricao_necessidade"].strip()

    elif secoes:
        partes_desc = []
        for chave in [
            "Contexto Institucional",
            "Diagnóstico da Situação Atual",
            "Fundamentação da Necessidade",
        ]:
            v = secoes.get(chave)
            if isinstance(v, str) and v.strip():
                partes_desc.append(v.strip())

        descricao_txt = "\n\n".join(partes_desc).strip()

    if not descricao_txt:
        descricao_txt = _to_str(campos.get("conteudo") or campos.get("descricao") or "")

    # ------------------------------------------------------------
    # MOTIVAÇÃO / OBJETIVOS / JUSTIFICATIVA
    # (Objetivos + Resultados + Benefícios + Justificativa + Riscos)
    # ------------------------------------------------------------
    motivacao_txt = ""

    if isinstance(campos.get("motivacao"), str) and campos["motivacao"].strip():
        motivacao_txt = campos["motivacao"].strip()

    elif secoes:
        partes_mot = []
        for chave in [
            "Objetivos da Contratação",
            "Resultados Esperados",
            "Benefícios Institucionais",
            "Justificativa Legal",
            "Riscos da Não Contratação",
        ]:
            v = secoes.get(chave)
            if isinstance(v, str) and v.strip():
                partes_mot.append(v.strip())

        motivacao_txt = "\n\n".join(partes_mot).strip()

    return {
        "unidade_demandante": unidade,
        "responsavel": responsavel,
        "prazo_estimado": prazo,
        "descricao": descricao_txt,
        "motivacao": motivacao_txt,
        "valor_estimado": valor_estimado,
    }


def _montar_texto_narrativo_inicial(
    dados_brutos: Dict[str, Any],
    secoes: Dict[str, str],
    campos_tradicionais: Dict[str, str],
) -> str:
    """
    Monta o texto_narrativo inicial. Se já existir no JSON, usa direto.
    Caso contrário, monta uma versão numerada a partir das seções,
    ou, em último caso, a partir de descrição + motivação.
    """
    existente = dados_brutos.get("texto_narrativo")
    if isinstance(existente, str) and existente.strip():
        return existente.strip()

    # Tentar construir com base nas 11 seções
    partes = []
    idx = 1
    for nome in SECOES_DFD:
        texto_secao = secoes.get(nome, "").strip()
        if texto_secao:
            partes.append(f"{idx}. {texto_secao}")
            idx += 1

    if partes:
        return "\n\n".join(partes)

    # Fallback: descrição + motivação
    descricao = campos_tradicionais.get("descricao", "").strip()
    motivacao = campos_tradicionais.get("motivacao", "").strip()

    partes = []
    if descricao:
        partes.append(f"1. {descricao}")
    if motivacao:
        partes.append(f"2. {motivacao}")

    return "\n\n".join(partes).strip()


# ======================================================================
# 1️⃣ Carregar dados já existentes (sessão ou arquivos)
# ======================================================================
dfd_dados = obter_dfd_da_sessao()

# ======================================================================
# ASSISTENTE IA – Ferramentas de automação
# ======================================================================
st.markdown('<div class="ia-block">', unsafe_allow_html=True)
st.markdown("### Assistente IA")
st.caption("Processamento automático: requer documentos no módulo Insumos")

col_ia1, col_ia2, col_ia3 = st.columns(3)

with col_ia1:
    if st.button("Gerar rascunho automático", use_container_width=True, type="primary", key="btn_ia_gerar"):
        try:
            with st.spinner("Processando documento..."):
                dfd_ai = gerar_rascunho_dfd_com_ia()

            if dfd_ai:
                st.success("Rascunho gerado com sucesso")
                st.rerun()
            else:
                st.warning("Nenhum insumo encontrado. Verifique o módulo Insumos.")
        except Exception as e:
            st.error(f"Erro ao processar: {e}")

with col_ia2:
    if st.button("Enviar para ETP", use_container_width=True, disabled=not dfd_dados, key="btn_enviar_etp"):
        try:
            import os
            from datetime import datetime
            
            base = os.path.join("exports", "insumos", "json")
            os.makedirs(base, exist_ok=True)
            
            payload = {
                "artefato": "ETP",
                "origem": "DFD_estruturado",
                "data_processamento": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "status": "ok",
                "campos_ai": dfd_dados,
                "conteudo_textual": dfd_dados.get("texto_narrativo", ""),
            }
            
            arq_ultimo = os.path.join(base, "ETP_ultimo.json")
            with open(arq_ultimo, "w", encoding="utf-8") as f:
                json.dump(payload, f, ensure_ascii=False, indent=2)
            
            st.success("Dados enviados para o módulo ETP")
            st.info("Acesse o módulo ETP para continuar")
            
        except Exception as e:
            st.error(f"Erro: {e}")

with col_ia3:
    st.write("")  # Espaçamento

st.markdown('</div>', unsafe_allow_html=True)

# ======================================================================
# REFINAMENTO ITERATIVO – Ajustes por seção
# ======================================================================
# Verificar se houve atualização via refinamento
dfd_dados = render_refinamento_iterativo(
    secoes_disponiveis=SECOES_DFD,
    dados_atuais=dfd_dados if dfd_dados else {},
    artefato="DFD",
    campos_simples=["unidade_demandante", "responsavel", "prazo_estimado", "valor_estimado",
                    "descricao_necessidade", "motivacao", "texto_narrativo"]
)

st.markdown("---")

# Se não há dados prévios, inicializa com estrutura vazia para permitir preenchimento manual
if not dfd_dados:
    st.info("Nenhum DFD encontrado. Opções disponíveis:")
    col1, col2 = st.columns(2)
    with col1:
        st.markdown("**Processar automaticamente:** Envie documentos no módulo Insumos e use o Assistente IA")
    with col2:
        st.markdown("**Preencher manualmente:** Use o formulário abaixo")
    st.markdown("---")
    
    # Inicializa estrutura vazia
    dfd_dados = {
        "unidade_demandante": "",
        "responsavel": "",
        "prazo_estimado": "",
        "valor_estimado": "0,00",
        "descricao_necessidade": "",
        "motivacao": "",
        "secoes": {secao: "" for secao in SECOES_DFD},
        "lacunas": [],
        "texto_narrativo": ""
    }

# Caso ainda venha algo como {"DFD": {...}}, normalizar
if isinstance(dfd_dados, dict) and "DFD" in dfd_dados:
    dfd_dados = dfd_dados.get("DFD") or {}

campos_trad = _extrair_admin_e_campos_tradicionais(dfd_dados)
secoes_iniciais = _extrair_secoes(dfd_dados)
lacunas_iniciais = _extrair_lacunas(dfd_dados)
texto_narrativo_inicial = _montar_texto_narrativo_inicial(
    dfd_dados,
    secoes_iniciais,
    campos_trad,
)

with st.expander("Visualizar dados importados (JSON)", expanded=False):
    st.json(dfd_dados)

# ======================================================================
# FORMULÁRIO DFD
# ======================================================================
st.subheader("Formulário DFD")

with st.form(key="form_dfd_moderno"):

    st.markdown("### Dados Administrativos")

    col1, col2 = st.columns(2)
    unidade = col1.text_input("Unidade Demandante", value=campos_trad["unidade_demandante"])
    responsavel = col2.text_input("Responsável", value=campos_trad["responsavel"])

    col3, col4 = st.columns(2)
    prazo = col3.text_input("Prazo Estimado", value=campos_trad["prazo_estimado"])
    valor_estimado = col4.text_input("Valor Estimado (R$)", value=campos_trad["valor_estimado"])

    st.markdown("---")
    st.markdown("### Síntese da Demanda")

    descricao = st.text_area(
        "Descrição da Necessidade",
        value=campos_trad["descricao"],
        height=180,
    )

    motivacao = st.text_area(
        "Motivação e Objetivos",
        value=campos_trad["motivacao"],
        height=180,
    )

    st.markdown("---")
    st.markdown("### Texto Narrativo Consolidado")

    texto_narrativo = st.text_area(
        "Texto completo estruturado",
        value=texto_narrativo_inicial,
        height=260,
    )

    st.markdown("---")
    st.markdown("### Seções Estruturadas (11 seções padrão)")

    secoes_editadas: Dict[str, str] = {}

    with st.expander("Editar seções individualmente", expanded=False):
        for nome_secao in SECOES_DFD:
            secoes_editadas[nome_secao] = st.text_area(
                nome_secao,
                value=secoes_iniciais.get(nome_secao, ""),
                height=140,
            )

    st.markdown("---")
    st.markdown("### Lacunas Identificadas")

    if lacunas_iniciais:
        for item in lacunas_iniciais:
            st.markdown(f"- {item}")
    else:
        st.caption("Nenhuma lacuna identificada")

    submit = st.form_submit_button("Salvar DFD")


# ======================================================================
# SALVAMENTO
# ======================================================================
if submit:
    dfd_final = {
        # Administrativos
        "unidade_demandante": unidade,
        "responsavel": responsavel,
        "prazo_estimado": prazo,
        "valor_estimado": valor_estimado,
        # Campo tradicional (síntese)
        "descricao_necessidade": descricao,
        "motivacao": motivacao,
        # Estrutura moderna
        "texto_narrativo": texto_narrativo,
        "secoes": secoes_editadas,
        "lacunas": lacunas_iniciais,
        # Metadado auxiliar opcional
        "atualizado_em": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "origem": "dfd_moderno_streamlit",
    }

    caminho = salvar_dfd_em_json(dfd_final, origem="formulario_dfd_moderno_streamlit")

    st.success("DFD salvo com sucesso")
    st.caption(f"Arquivo: `{caminho}`")
    st.json(dfd_final)

# ======================================================================
# EXPORTAÇÃO
# ======================================================================
st.subheader("Exportar Documento")

if st.button("Baixar DFD (DOCX)"):
    doc = Document()

    doc.add_heading("Formalização da Demanda (DFD)", level=1)

    # 1. Dados Administrativos
    doc.add_heading("1. Dados Administrativos", level=2)
    doc.add_paragraph(f"Unidade Demandante: {unidade}")
    doc.add_paragraph(f"Responsável pela Demanda: {responsavel}")
    doc.add_paragraph(f"Prazo Estimado: {prazo}")
    doc.add_paragraph(f"Estimativa de Valor: R$ {valor_estimado}")

    # 2. Texto narrativo consolidado
    doc.add_heading("2. Texto Narrativo Consolidado", level=2)
    doc.add_paragraph(texto_narrativo)

    # 3. Síntese tradicional
    doc.add_heading("3. Síntese Tradicional do DFD", level=2)
    doc.add_heading("3.1 Descrição da Necessidade", level=3)
    doc.add_paragraph(descricao)
    doc.add_heading("3.2 Motivação / Objetivos / Justificativa", level=3)
    doc.add_paragraph(motivacao)

    # 4. Seções estruturadas
    doc.add_heading("4. Seções Estruturadas (Modelo Moderno-Governança)", level=2)
    for nome_secao in SECOES_DFD:
        doc.add_heading(nome_secao, level=3)
        doc.add_paragraph(secoes_editadas.get(nome_secao, ""))

    # 5. Lacunas
    doc.add_heading("5. Lacunas Identificadas", level=2)
    if lacunas_iniciais:
        for item in lacunas_iniciais:
            doc.add_paragraph(f"- {item}")
    else:
        doc.add_paragraph("Não foram identificadas lacunas relevantes pela IA para este DFD.")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    st.download_button(
        label="⬇️ Download DOCX (DFD completo)",
        data=buffer,
        file_name="DFD_consolidado_moderno.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
