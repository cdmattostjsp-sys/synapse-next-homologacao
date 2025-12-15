import sys
from pathlib import Path
ROOT = Path(__file__).resolve().parents[2]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

# ==============================
# pages/06_📜 Edital – Minuta do Edital.py  –  SynapseNext / SAAB TJSP
# ==============================

import os, sys, json
from datetime import datetime
from io import BytesIO
from pathlib import Path
import streamlit as st
from docx import Document

# ==========================================================
# 🔍 Imports e configuração de ambiente
# ==========================================================

from utils.ui_components import aplicar_estilo_global, exibir_cabecalho_padrao
from home_utils.sidebar_organizer import apply_sidebar_grouping
from utils.integration_edital import integrar_com_contexto, gerar_edital_com_ia
from home_utils.refinamento_ia import render_refinamento_iterativo

st.set_page_config(page_title="📜 Edital – Minuta", layout="wide", page_icon="📜")
apply_sidebar_grouping()

# Estilo institucional PJe-inspired
st.markdown("""
<style>
/* ============================================
   PADRÃO VISUAL PJe-INSPIRED - SYNAPSE NEXT
   Versão: 2025.1-homolog
   ============================================ */

/* Título principal - tamanho reduzido para sobriedade */
h1 {
    font-size: 1.8rem !important;
    font-weight: 500 !important;
    color: #2c3e50 !important;
    margin-bottom: 0.3rem !important;
}

/* Caption institucional */
.caption {
    color: #6c757d;
    font-size: 0.9rem;
    margin-bottom: 1rem;
}

/* Bloco de IA - destaque sutil */
.ia-block {
    border: 1px solid #d0d7de;
    border-radius: 3px;
    padding: 1rem 1.2rem;
    background-color: #f0f2f5;
    margin: 1rem 0 1.2rem 0;
}
.ia-block h3 {
    font-size: 1rem;
    font-weight: 600;
    color: #1f2937;
    margin: 0 0 0.6rem 0;
    letter-spacing: -0.01em;
}

/* Seções com fundo cinza - contraste melhorado */
h3 {
    font-size: 1.1rem !important;
    font-weight: 500 !important;
    color: #374151 !important;
    background-color: #e5e7eb !important;
    padding: 0.6rem 0.8rem !important;
    border-radius: 3px !important;
    margin-top: 1.5rem !important;
    margin-bottom: 1rem !important;
}

/* Botões - destaque apenas para ações principais */
div.stButton > button {
    border-radius: 3px;
    font-weight: 500;
    border: 1px solid #d0d7de;
}
div.stButton > button[kind="primary"] {
    background-color: #0969da !important;
    border-color: #0969da !important;
}

/* Formulário clean */
.stTextInput label, .stTextArea label {
    font-weight: 500;
    color: #1f2937;
    font-size: 0.9rem;
}

/* Expander refinamento com destaque discreto */
details {
    border: 1px solid #d0d7de;
    border-radius: 3px;
    padding: 0.5rem;
    background-color: #ffffff;
}
summary {
    font-weight: 500;
    color: #0969da;
    cursor: pointer;
}
</style>
""", unsafe_allow_html=True)

# ==========================================================
# 🏛️ Cabeçalho institucional
# ==========================================================
st.markdown("<h1>Minuta do Edital de Licitação</h1>", unsafe_allow_html=True)
st.markdown("<p class='caption'>Geração automatizada com IA institucional a partir dos artefatos DFD, ETP e TR</p>", unsafe_allow_html=True)
st.divider()

# ==========================================================
# 🔍 Carregamento inteligente: INSUMOS + Contexto → EditalAgent → Formulário
# ==========================================================

# Paths dos arquivos
INSUMO_EDITAL_PATH = os.path.join("exports", "insumos", "json", "EDITAL_ultimo.json")
EDITAL_JSON_PATH = os.path.join("exports", "edital_data.json")

# ==========================================================
# 📝 Definição dos campos do Edital para refinamento iterativo
# ==========================================================
CAMPOS_EDITAL = [
    "numero_edital",
    "data_publicacao",
    "objeto",
    "tipo_licitacao",
    "criterio_julgamento",
    "condicoes_participacao",
    "exigencias_habilitacao",
    "obrigacoes_contratada",
    "prazo_execucao",
    "fontes_recursos",
    "gestor_fiscal",
    "observacoes_gerais"
]

# Integrar contexto de DFD/ETP/TR
contexto = integrar_com_contexto(st.session_state)
if contexto:
    st.success(f"📎 Contexto integrado: {', '.join(contexto.keys())}")

edital_salvo = {}
insumo_detectado = False

# 1️⃣ Verificar se existe Edital processado
if os.path.exists(EDITAL_JSON_PATH):
    try:
        with open(EDITAL_JSON_PATH, "r", encoding="utf-8") as f:
            dados_edital = json.load(f)
            edital_salvo = dados_edital.get("EDITAL", {})
            if edital_salvo and any(v for v in edital_salvo.values() if v):
                st.success("📎 Edital processado carregado (exports/edital_data.json)")
    except Exception as e:
        st.warning(f"⚠️ Erro ao carregar Edital: {e}")

# 2️⃣ Se não houver Edital processado, verificar INSUMO bruto
if not edital_salvo or not any(v for v in edital_salvo.values() if v):
    if os.path.exists(INSUMO_EDITAL_PATH):
        insumo_detectado = True
        st.info("📄 Insumo EDITAL detectado. Use o botão **'Processar com IA'** para extrair os 12 campos automaticamente.")
        
        # Carregar texto bruto do insumo (para preview)
        try:
            with open(INSUMO_EDITAL_PATH, "r", encoding="utf-8") as f:
                insumo_data = json.load(f)
                texto_bruto = insumo_data.get("conteudo_textual", "")
                
                if texto_bruto and len(texto_bruto) > 100:
                    with st.expander("👁️ Preview do insumo carregado", expanded=False):
                        st.text_area(
                            "Texto extraído do PDF/DOCX:",
                            texto_bruto[:1000] + "..." if len(texto_bruto) > 1000 else texto_bruto,
                            height=200,
                            disabled=True
                        )
        except Exception as e:
            st.warning(f"⚠️ Erro ao ler insumo: {e}")

# 3️⃣ Nenhum dado disponível
if not edital_salvo and not insumo_detectado and not contexto:
    st.info("ℹ️ Nenhum Edital ou insumo detectado. Faça upload no módulo **🔧 Insumos** ou preencha dados em DFD/ETP/TR primeiro.")

# ==========================================================
# 🧾 Formulário Edital – 12 Campos Estruturados
# ==========================================================
st.markdown("### Assistente IA")
st.caption("Processamento automático: requer insumos do módulo anterior")

col_ia1, col_ia2, col_ia3 = st.columns(3)

with col_ia1:
    if st.button("⚡ Processar com IA", use_container_width=True, type="primary", key="btn_ia_processar"):
        try:
            with st.spinner("Processando Edital com IA especializada..."):
                resultado = gerar_edital_com_ia(contexto_previo=contexto)
                
            if "erro" in resultado:
                st.error(f"❌ {resultado['erro']}")
            else:
                st.success("Edital estruturado com sucesso")
                
                # Exibir resumo
                edital_processado = resultado.get("EDITAL", {})
                campos_ia = sum(1 for v in edital_processado.values() if v and v.strip())
                
                col_a, col_b, col_c = st.columns(3)
                with col_a:
                    st.metric("Campos", f"{campos_ia}/12")
                with col_b:
                    st.metric("Número", edital_processado.get("numero_edital", "N/A")[:20])
                with col_c:
                    st.metric("Modalidade", edital_processado.get("tipo_licitacao", "N/A")[:20])
                
                # Salvar dados para integração com Validador
                st.session_state["edital_campos_ai"] = edital_processado
                st.session_state["edital_processado_agora"] = True
                
                st.success("✅ Edital disponível para validação no módulo **🧩 Validador de Editais**")
                st.rerun()
        except Exception as e:
            st.error(f"Erro ao processar: {e}")

with col_ia2:
    if st.button("📤 Enviar para Contrato", use_container_width=True, disabled=not edital_salvo, key="btn_enviar_contrato"):
        try:
            from datetime import datetime
            
            base = os.path.join("exports", "insumos", "json")
            os.makedirs(base, exist_ok=True)
            
            payload = {
                "artefato": "CONTRATO",
                "origem": "EDITAL_estruturado",
                "data_processamento": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "status": "ok",
                "campos_ai": edital_salvo,
                "conteudo_textual": "",
            }
            
            arq_ultimo = os.path.join(base, "CONTRATO_ultimo.json")
            with open(arq_ultimo, "w", encoding="utf-8") as f:
                json.dump(payload, f, ensure_ascii=False, indent=2)
            
            st.success("Dados enviados para o módulo Contrato")
            st.info("Acesse o módulo Contrato para continuar")
            
        except Exception as e:
            st.error(f"Erro: {e}")

with col_ia3:
    st.write("")  # Espaçamento

st.divider()

# Refinamento iterativo
campos_simples = ["numero_edital", "data_publicacao", "tipo_licitacao", "criterio_julgamento", 
                  "prazo_execucao", "fontes_recursos", "gestor_fiscal"]

edital_salvo = render_refinamento_iterativo(
    secoes_disponiveis=CAMPOS_EDITAL,
    dados_atuais=edital_salvo if edital_salvo else {},
    artefato="EDITAL",
    campos_simples=campos_simples
)

st.divider()

st.markdown("### 📋 Identificação")
col1, col2 = st.columns(2)
with col1:
    numero_edital = st.text_input(
        "Número do Edital",
        value=edital_salvo.get("numero_edital", ""),
        key="ed_numero",
        help="Identificação única do edital"
    )
    
with col2:
    data_publicacao = st.text_input(
        "Data de Publicação",
        value=edital_salvo.get("data_publicacao", ""),
        key="ed_data",
        help="Data de divulgação do edital"
    )

st.markdown("### 📋 Objeto e Modalidade")

objeto = st.text_area(
    "1. Objeto da Licitação",
    value=edital_salvo.get("objeto", ""),
    height=120,
    key="ed_objeto",
    help="Descrição do que será licitado"
)

col3, col4 = st.columns(2)
with col3:
    tipo_licitacao = st.text_input(
        "2. Tipo de Licitação",
        value=edital_salvo.get("tipo_licitacao", ""),
        key="ed_tipo",
        help="Pregão eletrônico, concorrência, etc."
    )
    
with col4:
    criterio_julgamento = st.text_input(
        "3. Critério de Julgamento",
        value=edital_salvo.get("criterio_julgamento", ""),
        key="ed_criterio",
        help="Menor preço, melhor técnica, etc."
    )

st.markdown("### Requisitos e Condições")

# Seção 4: Condições de Participação (coluna única)
condicoes_participacao = st.text_area(
    "4. Condições de Participação",
    value=edital_salvo.get("condicoes_participacao", ""),
    height=120,
    key="ed_cond",
    help="Requisitos para participar"
)

# Seção 5: Exigências de Habilitação (coluna única)
exigencias_habilitacao = st.text_area(
    "5. Exigências de Habilitação",
    value=edital_salvo.get("exigencias_habilitacao", ""),
    height=120,
    key="ed_exig",
    help="Documentação necessária"
)

# Seção 6: Obrigações da Contratada (coluna única)
obrigacoes_contratada = st.text_area(
    "6. Obrigações da Contratada",
    value=edital_salvo.get("obrigacoes_contratada", ""),
    height=120,
    key="ed_obrig",
    help="Deveres da empresa vencedora"
)

# Seção 10: Observações Gerais (coluna única)
observacoes_gerais = st.text_area(
    "10. Observações Gerais",
    value=edital_salvo.get("observacoes_gerais", ""),
    height=120,
    key="ed_obs",
    help="Informações complementares"
)

st.markdown("### Informações Administrativas")

col7, col8, col9 = st.columns(3)
with col7:
    prazo_execucao = st.text_input(
        "7. Prazo de Execução",
        value=edital_salvo.get("prazo_execucao", ""),
        key="ed_prazo"
    )
    
with col8:
    fontes_recursos = st.text_input(
        "8. Fontes de Recursos",
        value=edital_salvo.get("fontes_recursos", ""),
        key="ed_fonte"
    )
    
with col9:
    gestor_fiscal = st.text_input(
        "9. Gestor/Fiscal",
        value=edital_salvo.get("gestor_fiscal", ""),
        key="ed_gestor"
    )

st.divider()

# Botões de ação
col_salvar, col_baixar = st.columns(2)

with col_salvar:
    if st.button("Salvar Edital", type="secondary", use_container_width=True):
        edital_completo = {
            "numero_edital": numero_edital,
            "data_publicacao": data_publicacao,
            "objeto": objeto,
            "tipo_licitacao": tipo_licitacao,
            "criterio_julgamento": criterio_julgamento,
            "condicoes_participacao": condicoes_participacao,
            "exigencias_habilitacao": exigencias_habilitacao,
            "obrigacoes_contratada": obrigacoes_contratada,
            "prazo_execucao": prazo_execucao,
            "fontes_recursos": fontes_recursos,
            "gestor_fiscal": gestor_fiscal,
            "observacoes_gerais": observacoes_gerais,
        }
        
        try:
            os.makedirs("exports", exist_ok=True)
            with open(EDITAL_JSON_PATH, "w", encoding="utf-8") as f:
                json.dump({"EDITAL": edital_completo}, f, ensure_ascii=False, indent=2)
            st.success("Edital salvo com sucesso")
        except Exception as e:
            st.error(f"Erro ao salvar: {e}")

with col_baixar:
    if st.button("Baixar Edital (DOCX)", use_container_width=True):
        doc = Document()
        
        doc.add_heading("Minuta do Edital de Licitação", level=1)
        doc.add_paragraph("Lei 14.133/2021 - Nova Lei de Licitações")
        
        # Identificação
        doc.add_heading("Identificação", level=2)
        doc.add_paragraph(f"Número do Edital: {numero_edital if numero_edital else '[Não preenchido]'}")
        doc.add_paragraph(f"Data de Publicação: {data_publicacao if data_publicacao else '[Não preenchido]'}")
        
        # Objeto e Modalidade
        doc.add_heading("Objeto e Modalidade", level=2)
        doc.add_heading("1. Objeto da Licitação", level=3)
        doc.add_paragraph(objeto if objeto and objeto.strip() else "[Não preenchido]")
        doc.add_paragraph(f"2. Tipo de Licitação: {tipo_licitacao if tipo_licitacao else '[Não preenchido]'}")
        doc.add_paragraph(f"3. Critério de Julgamento: {criterio_julgamento if criterio_julgamento else '[Não preenchido]'}")
        
        # Requisitos e Condições
        doc.add_heading("Requisitos e Condições", level=2)
        
        secoes_nomes = [
            ("4. Condições de Participação", condicoes_participacao),
            ("5. Exigências de Habilitação", exigencias_habilitacao),
            ("6. Obrigações da Contratada", obrigacoes_contratada),
            ("10. Observações Gerais", observacoes_gerais),
        ]
        
        for nome, conteudo in secoes_nomes:
            doc.add_heading(nome, level=3)
            if conteudo and conteudo.strip():
                doc.add_paragraph(conteudo)
            else:
                doc.add_paragraph("[Não preenchido]")
        
        # Informações Administrativas
        doc.add_heading("Informações Administrativas", level=2)
        doc.add_paragraph(f"7. Prazo de Execução: {prazo_execucao if prazo_execucao else '[Não preenchido]'}")
        doc.add_paragraph(f"8. Fontes de Recursos: {fontes_recursos if fontes_recursos else '[Não preenchido]'}")
        doc.add_paragraph(f"9. Gestor/Fiscal: {gestor_fiscal if gestor_fiscal else '[Não preenchido]'}")
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        st.download_button(
            label="⬇️ Download DOCX (Edital completo)",
            data=buffer,
            file_name="Edital_Minuta_Lei14133.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )

st.caption("Dica: Use o Assistente IA para preencher automaticamente os campos a partir dos insumos carregados")
