# ==========================================================
# utils/integration_edital.py
# SynapseNext – SAAB / TJSP
# Integração estável do módulo EDITAL (modo híbrido)
# ==========================================================

from __future__ import annotations

import os
import io
import re
import json
from pathlib import Path
from datetime import datetime
from typing import Dict, Any, Optional

# Streamlit é opcional para permitir testes sem UI
try:
    import streamlit as st  # type: ignore
except Exception:  # pragma: no cover
    st = None

# -----------------------------
# Dependências opcionais de extração
# -----------------------------
try:
    import fitz  # PyMuPDF
except Exception:  # pragma: no cover
    fitz = None

try:
    import docx2txt
except Exception:  # pragma: no cover
    docx2txt = None

try:
    from docx import Document
except Exception:  # pragma: no cover
    Document = None

# -----------------------------
# OpenAI – criação tardia (opcional)
# -----------------------------
def _get_openai_client():
    api_key = None
    if st and getattr(st, "secrets", None):
        api_key = st.secrets.get("OPENAI_API_KEY")
    api_key = api_key or os.getenv("OPENAI_API_KEY") or os.getenv("openai_api_key")

    try:
        from openai import OpenAI  # openai>=1.x
    except Exception:
        return None, None

    if not api_key:
        return None, None

    try:
        return OpenAI(api_key=api_key), "gpt-4o-mini"
    except Exception:
        return None, None


# -----------------------------
# Paths institucionais (opcional)
# -----------------------------
BASE_DIR = Path(__file__).resolve().parents[1] if "__file__" in globals() else Path.cwd()
EXPORTS_DIR = BASE_DIR / "exports"
EXPORTS_DIR.mkdir(parents=True, exist_ok=True)

KB_EDITAL_DIR = (BASE_DIR / "knowledge_base" / "edital")  # tolerante: não obrigatório


# ==========================================================
# 📚 Leitura de modelos da KB (tolerante)
# ==========================================================
def ler_modelos_edital() -> str:
    textos = []
    if KB_EDITAL_DIR.exists():
        for arq in KB_EDITAL_DIR.glob("*.txt"):
            try:
                textos.append(arq.read_text(encoding="utf-8"))
            except Exception:
                pass
    return "\n\n".join(textos)


# ==========================================================
# 🔗 Integração de contexto (DFD, ETP, TR) – modo híbrido
# ==========================================================
def integrar_com_contexto(session_state: dict | None) -> dict:
    """
    Constrói um contexto integrado com o que existir na sessão:
    DFD + ETP + TR. É tolerante à ausência de qualquer um deles.
    """
    ss = session_state or {}
    contexto = {}
    for chave in ("dfd_campos_ai", "etp_campos_ai", "tr_campos_ai"):
        bloco = ss.get(chave)
        if isinstance(bloco, dict) and bloco:
            contexto[chave] = bloco
    return contexto

# Alias de compatibilidade
def consolidar_contexto(session_state: dict | None) -> dict:
    return integrar_com_contexto(session_state)


# ==========================================================
# 🧾 Extração e limpeza de texto do insumo
# ==========================================================
def extrair_texto_arquivo(arquivo) -> str:
    nome = getattr(arquivo, "name", "").lower()

    def limpar(txt: str) -> str:
        txt = re.sub(r"\s+", " ", txt or "")
        txt = re.sub(r"[^\w\s.,;:!?()/%\-–—ºª°]", "", txt)
        return txt.strip()

    try:
        # PDF
        if nome.endswith(".pdf"):
            data = arquivo.read()
            try:
                arquivo.seek(0)
            except Exception:
                pass
            if fitz is not None:
                try:
                    texto = ""
                    with fitz.open(stream=data, filetype="pdf") as pdf:
                        for p in pdf:
                            texto += p.get_text("text") + "\n"
                    return limpar(texto)
                except Exception:
                    pass
            return ""  # sem fallback obrigatório p/ PDF (evita hard deps)

        # DOCX
        if nome.endswith(".docx") and docx2txt is not None:
            data = arquivo.read()
            try:
                arquivo.seek(0)
            except Exception:
                pass
            try:
                return limpar(docx2txt.process(io.BytesIO(data)))
            except Exception:
                pass

        # TXT
        if nome.endswith(".txt"):
            data = arquivo.read()
            try:
                arquivo.seek(0)
            except Exception:
                pass
            try:
                return limpar(data.decode("utf-8", errors="ignore"))
            except Exception:
                return limpar(data.decode("latin-1", errors="ignore"))
    except Exception:
        pass
    return ""


# ==========================================================
# 🧠 IA opcional para estruturar campos do edital
# ==========================================================
def _chamar_ia_edital(texto_insumo: str, modelos: str, contexto: dict) -> Dict[str, Any]:
    client, model = _get_openai_client()
    if client is None or not model or not texto_insumo.strip():
        return {}

    system_prompt = (
        "Você é um agente institucional do TJSP (SAAB) especializado em minutas de EDITAL (Lei 14.133/2021). "
        "Use linguagem padrão SAAB/TJSP. Retorne somente JSON."
    )

    user_prompt = f"""
Texto do insumo (base):
\"\"\"{texto_insumo[:10000]}\"\"\"

Contexto cumulativo disponível (DFD, ETP, TR):
\"\"\"{json.dumps(contexto or {}, ensure_ascii=False)}\"\"\"

Modelos institucionais (KB):
\"\"\"{(modelos or '')[:8000]}\"\"\"

Retorne APENAS um JSON com os campos:
{{
  "objeto": "",
  "tipo_licitacao": "",
  "criterio_julgamento": "",
  "condicoes_participacao": "",
  "exigencias_habilitacao": "",
  "obrigacoes_contratada": "",
  "prazo_execucao": "",
  "fontes_recursos": "",
  "gestor_fiscal": "",
  "observacoes_gerais": ""
}}
"""

    try:
        resp = client.chat.completions.create(
            model=model,
            messages=[{"role": "system", "content": system_prompt},
                      {"role": "user", "content": user_prompt}],
            temperature=0.3,
            max_tokens=900,
        )
        conteudo = (resp.choices[0].message.content or "").strip()
        m = re.search(r"\{.*\}", conteudo, re.S)
        return json.loads(m.group(0)) if m else {}
    except Exception:
        return {}


# ==========================================================
# 🧹 Normalização e defaults (modo híbrido)
# ==========================================================
def _normalizar_campos(campos: Dict[str, Any], contexto: dict) -> Dict[str, str]:
    hoje = datetime.now()
    numero_auto = f"TJSP-PE-{hoje.year}-{hoje.strftime('%m')}{hoje.strftime('%d')}"
    data_auto = hoje.strftime("%d/%m/%Y")

    dfd = (contexto or {}).get("dfd_campos_ai", {})
    etp = (contexto or {}).get("etp_campos_ai", {})
    tr  = (contexto or {}).get("tr_campos_ai", {})

    def first_nonempty(*vals):
        for v in vals:
            if isinstance(v, str) and v.strip():
                return v
        return ""

    defaults = {
        "objeto": first_nonempty(tr.get("objeto"), etp.get("objeto"), dfd.get("objeto")),
        "tipo_licitacao": "Pregão eletrônico",
        "criterio_julgamento": first_nonempty(tr.get("criterios_de_julgamento"), "Menor preço global"),
        "condicoes_participacao": "",
        "exigencias_habilitacao": "",
        "obrigacoes_contratada": first_nonempty(tr.get("obrigacoes_da_contratada")),
        "prazo_execucao": first_nonempty(tr.get("prazo_execucao")),
        "fontes_recursos": first_nonempty(tr.get("fonte_recurso")),
        "gestor_fiscal": first_nonempty(dfd.get("responsavel"), dfd.get("responsavel_tecnico")),
        "observacoes_gerais": "",
        "numero_edital": numero_auto,
        "data_publicacao": data_auto,
        "unidade_solicitante": first_nonempty(dfd.get("unidade_solicitante")),
        "responsavel": first_nonempty(dfd.get("responsavel"), dfd.get("responsavel_tecnico")),
    }

    result = {}
    for k, v in defaults.items():
        result[k] = (campos.get(k) or v or "").strip()
        result[k] = re.sub(r"\s+", " ", result[k]).strip()

    return result

# Alias compatível com páginas antigas
def normalizar_campos_edital(campos: Dict[str, Any], contexto: dict) -> Dict[str, str]:
    return _normalizar_campos(campos, contexto)


# ==========================================================
# 🧱 Geração de rascunho e DOCX
# ==========================================================
def gerar_rascunho_edital(campos: Dict[str, str], modelos_referencia: str = "") -> str:
    linhas = [
        f"EDITAL Nº {campos.get('numero_edital','')}",
        f"Data de Publicação: {campos.get('data_publicacao','')}",
        "",
        f"Unidade Solicitante: {campos.get('unidade_solicitante','')}",
        f"Responsável: {campos.get('responsavel','')}",
        "",
        "1. DO OBJETO",
        campos.get("objeto",""),
        "",
        "2. DO TIPO E CRITÉRIO DE JULGAMENTO",
        f"Tipo de licitação: {campos.get('tipo_licitacao','')}",
        f"Critério de julgamento: {campos.get('criterio_julgamento','')}",
        "",
        "3. DAS CONDIÇÕES DE PARTICIPAÇÃO",
        campos.get("condicoes_participacao",""),
        "",
        "4. DAS EXIGÊNCIAS DE HABILITAÇÃO",
        campos.get("exigencias_habilitacao",""),
        "",
        "5. DAS OBRIGAÇÕES DA CONTRATADA",
        campos.get("obrigacoes_contratada",""),
        "",
        "6. DO PRAZO DE EXECUÇÃO",
        campos.get("prazo_execucao",""),
        "",
        "7. DAS FONTES DE RECURSOS",
        campos.get("fontes_recursos",""),
        "",
        "8. DO GESTOR/FISCAL DO CONTRATO",
        campos.get("gestor_fiscal",""),
        "",
        "9. DAS DISPOSIÇÕES FINAIS",
        campos.get("observacoes_gerais",""),
    ]
    if modelos_referencia:
        linhas += ["", "ANEXO – ORIENTAÇÕES INSTITUCIONAIS (KB)", (modelos_referencia[:1200] + " …")]
    return "\n".join(linhas)


def gerar_edital_docx(campos: Dict[str, str], texto_completo: Optional[str] = None) -> Optional[str]:
    if Document is None:
        return None  # evita falha se python-docx não estiver instalado

    doc = Document()
    doc.add_heading(f"EDITAL Nº {campos.get('numero_edital','')}", level=1)
    doc.add_paragraph(f"Data de Publicação: {campos.get('data_publicacao','')}")
    doc.add_paragraph(f"Unidade Solicitante: {campos.get('unidade_solicitante','')}")
    doc.add_paragraph(f"Responsável: {campos.get('responsavel','')}")
    doc.add_paragraph("")

    def bloco(titulo: str, corpo: str):
        doc.add_heading(titulo, level=2)
        doc.add_paragraph(corpo if corpo else "—")

    bloco("1. DO OBJETO", campos.get("objeto",""))
    bloco("2. DO TIPO E CRITÉRIO DE JULGAMENTO",
          f"Tipo: {campos.get('tipo_licitacao','')}. Critério: {campos.get('criterio_julgamento','')}.")
    bloco("3. DAS CONDIÇÕES DE PARTICIPAÇÃO", campos.get("condicoes_participacao",""))
    bloco("4. DAS EXIGÊNCIAS DE HABILITAÇÃO", campos.get("exigencias_habilitacao",""))
    bloco("5. DAS OBRIGAÇÕES DA CONTRATADA", campos.get("obrigacoes_contratada",""))
    bloco("6. DO PRAZO DE EXECUÇÃO", campos.get("prazo_execucao",""))
    bloco("7. DAS FONTES DE RECURSOS", campos.get("fontes_recursos",""))
    bloco("8. DO GESTOR/FISCAL DO CONTRATO", campos.get("gestor_fiscal",""))
    bloco("9. DAS DISPOSIÇÕES FINAIS", campos.get("observacoes_gerais",""))

    if texto_completo:
        doc.add_page_break()
        doc.add_heading("ANEXO – RASCUNHO INTEGRAL", level=2)
        for par in texto_completo.split("\n\n"):
            doc.add_paragraph(par)

    nome_arquivo = f"Edital_{campos.get('numero_edital','TJSP-PE')}.docx"
    caminho = str(EXPORTS_DIR / nome_arquivo)
    doc.save(caminho)
    return caminho


# ==========================================================
# 🚀 Função principal (entrada unificada para a página)
# ==========================================================
def processar_insumo_edital(arquivo, contexto_previo: dict | None = None, artefato: str = "EDITAL") -> dict:
    """
    1) Extrai texto do insumo (PDF/DOCX/TXT).
    2) Integra contexto (DFD/ETP/TR) – modo híbrido.
    3) Chama IA (se disponível) para estruturar campos.
    4) Normaliza e gera rascunho + DOCX.
    5) Retorna dicionário padronizado para pré-preenchimento.
    """
    texto = extrair_texto_arquivo(arquivo)
    if not texto:
        return {"erro": "Falha na extração de texto do insumo de EDITAL."}

    contexto = contexto_previo or integrar_com_contexto(st.session_state if st else None)
    modelos = ler_modelos_edital()
    campos_ia = _chamar_ia_edital(texto, modelos, contexto)
    campos = _normalizar_campos(campos_ia if isinstance(campos_ia, dict) else {}, contexto)

    rascunho = gerar_rascunho_edital(campos, modelos_referencia="")
    docx_path = gerar_edital_docx(campos, texto_completo=rascunho)

    payload = {
        "artefato": artefato,
        "nome_arquivo": getattr(arquivo, "name", ""),
        "status": "processado",
        "campos_ai": campos,
        "docx_path": docx_path,
        "contexto_usado": list((contexto or {}).keys()),
    }

    if st is not None:
        st.session_state["last_edital"] = payload

    return payload


# ==========================================================
# 🔁 Wrappers de compatibilidade (páginas antigas)
# ==========================================================
def processar_edital_dinamico(arquivo, contexto_previo: dict | None = None, artefato: str = "EDITAL") -> dict:
    """Alias histórico usado por algumas páginas."""
    return processar_insumo_edital(arquivo, contexto_previo=contexto_previo, artefato=artefato)
