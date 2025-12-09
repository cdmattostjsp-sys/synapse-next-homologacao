import sys
from pathlib import Path
ROOT = Path(__file__).resolve().parents[2]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

# ==========================================================
# pages/08_📜 Contrato.py – SynapseNext / SAAB TJSP
# ==========================================================
# Módulo final da jornada de contratação pública.
# Gera a minuta do contrato a partir de insumos cumulativos
# (DFD, ETP, TR, Edital) e processamento IA institucional.
# Refatorado para usar backend com lazy loading.
# ==========================================================

import os
from io import BytesIO
from datetime import datetime
import streamlit as st
from docx import Document
from pathlib import Path

from utils.ui_components import aplicar_estilo_global, exibir_cabecalho_padrao
from utils.integration_contrato import (
    processar_insumo_contrato,
    export_contrato_to_json,
    load_contrato_from_json,
    integrar_com_contexto,
)

# ==========================================================
# ⚙️ Configuração básica
# ==========================================================
st.set_page_config(page_title="📃 Contrato", layout="wide", page_icon="📃")
aplicar_estilo_global()

# ==========================================================
# 📥 Carregamento de dados persistidos (JSON)
# ==========================================================
dados_contrato_anterior = load_contrato_from_json()
if dados_contrato_anterior and "campos_ai" in dados_contrato_anterior:
    st.session_state["contrato_campos_ai"] = dados_contrato_anterior["campos_ai"]

# ==========================================================
# 🏛️ Cabeçalho institucional
# ==========================================================
exibir_cabecalho_padrao(
    "📃 Minuta do Contrato Administrativo",
    "Consolidação final dos artefatos da jornada de contratação (DFD → ETP → TR → Edital → Contrato)"
)
st.divider()

# ==========================================================
# 🔗 Dados cumulativos disponíveis
# ==========================================================
defaults = {}
for chave in ["dfd_campos_ai", "etp_campos_ai", "tr_campos_ai", "edital_campos_ai", "contrato_campos_ai"]:
    if chave in st.session_state:
        defaults.update(st.session_state[chave])

if defaults:
    st.success("📎 Dados recebidos automaticamente dos módulos anteriores (DFD, ETP, TR, Edital).")
else:
    st.info("Nenhum insumo ativo detectado. Você pode preencher manualmente ou aguardar integração via módulo INSUMOS.")

# ==========================================================
# 📤 Upload de insumo (opcional)
# ==========================================================
st.subheader("📤 Upload de Insumo (opcional)")
arquivo_upload = st.file_uploader(
    "Envie um arquivo de referência (PDF, DOCX, TXT) para processar com o backend:",
    type=["pdf", "docx", "txt"],
    help="O backend irá processar este arquivo e preencher automaticamente os campos abaixo."
)

if arquivo_upload is not None:
    if st.button("🔄 Processar Insumo com Backend"):
        with st.spinner("Processando insumo com backend integrado (lazy loading)..."):
            try:
                resultado_backend = processar_insumo_contrato(arquivo_upload)
                
                if resultado_backend["status"] == "processado":
                    # Salvar JSON
                    export_contrato_to_json(resultado_backend)
                    
                    # Atualizar session_state
                    st.session_state["contrato_campos_ai"] = resultado_backend["campos_ai"]
                    st.session_state["contrato_contexto"] = integrar_com_contexto(st.session_state)
                    
                    st.success("✅ Insumo processado com sucesso! Os campos abaixo foram preenchidos automaticamente.")
                    st.info(f"📄 Arquivo processado: {resultado_backend.get('nome_arquivo', 'N/A')}")
                    
                    # Recarregar página para mostrar dados atualizados
                    st.rerun()
                else:
                    st.warning("⚠️ Processamento concluído, mas status inesperado. Verifique os logs.")
                    
            except Exception as e:
                st.error(f"❌ Erro ao processar insumo: {e}")

st.divider()

# ==========================================================
# 🧾 Formulário – Campos contratuais
# ==========================================================
st.subheader("📄 Dados do Contrato")

col1, col2 = st.columns(2)
with col1:
    objeto = st.text_area("Objeto do Contrato", value=defaults.get("objeto", ""), height=100)
    partes = st.text_area("Partes Contratantes", value=defaults.get("partes", ""), height=80)
    vigencia = st.text_input("Vigência", value=defaults.get("vigencia", "12 meses a contar da assinatura"))
    valor_global = st.text_input("Valor Global", value=defaults.get("valor_global", ""))
    reajuste = st.text_area("Reajuste", value=defaults.get("reajuste", "Conforme índice oficial e cláusulas legais"), height=70)
    garantias = st.text_area("Garantias", value=defaults.get("garantias", ""), height=70)

with col2:
    prazos_pagamento = st.text_area("Prazos e Forma de Pagamento", value=defaults.get("prazos_pagamento", ""), height=70)
    obrigacoes_contratada = st.text_area("Obrigações da Contratada", value=defaults.get("obrigacoes_contratada", ""), height=100)
    obrigacoes_contratante = st.text_area("Obrigações da Contratante", value=defaults.get("obrigacoes_contratante", ""), height=100)
    fiscalizacao = st.text_area("Fiscalização e Acompanhamento", value=defaults.get("fiscalizacao", ""), height=70)
    penalidades = st.text_area("Penalidades", value=defaults.get("penalidades", ""), height=80)
    rescisao = st.text_area("Rescisão Contratual", value=defaults.get("rescisao", ""), height=80)
    foro = st.text_input("Foro Competente", value=defaults.get("foro", "Comarca de São Paulo/SP"))

st.divider()
observacoes_finais = st.text_area("Observações Finais", value=defaults.get("observacoes_finais", ""), height=70)

# ==========================================================
# 💾 Salvar manualmente campos editados
# ==========================================================
st.divider()
if st.button("💾 Salvar campos editados manualmente"):
    campos_manuais = {
        "objeto": objeto,
        "partes": partes,
        "vigencia": vigencia,
        "valor_global": valor_global,
        "reajuste": reajuste,
        "garantias": garantias,
        "prazos_pagamento": prazos_pagamento,
        "obrigacoes_contratada": obrigacoes_contratada,
        "obrigacoes_contratante": obrigacoes_contratante,
        "fiscalizacao": fiscalizacao,
        "penalidades": penalidades,
        "rescisao": rescisao,
        "foro": foro,
        "observacoes_finais": observacoes_finais
    }
    
    resultado_manual = {
        "artefato": "",
        "nome_arquivo": "edicao_manual",
        "status": "editado_manualmente",
        "campos_ai": campos_manuais
    }
    
    export_contrato_to_json(resultado_manual)
    st.session_state["contrato_campos_ai"] = campos_manuais
    st.session_state["contrato_contexto"] = integrar_com_contexto(st.session_state)
    st.success("✅ Campos salvos com sucesso!")

# ==========================================================
# 📄 Geração DOCX final
# ==========================================================
st.divider()
st.subheader("📄 Exportação da Minuta Contratual")

if st.button("📤 Gerar DOCX da Minuta Contratual"):
    with st.spinner("Gerando documento DOCX..."):
        doc = Document()
        doc.add_heading("MINUTA DO CONTRATO ADMINISTRATIVO", level=1)
        doc.add_heading("TJSP - Tribunal de Justiça de São Paulo", level=2)
        doc.add_paragraph("")
        
        doc.add_heading("1. OBJETO DO CONTRATO", level=2)
        doc.add_paragraph(objeto or "Não especificado")
        
        doc.add_heading("2. PARTES CONTRATANTES", level=2)
        doc.add_paragraph(partes or "Não especificado")
        
        doc.add_heading("3. VIGÊNCIA", level=2)
        doc.add_paragraph(vigencia or "Não especificado")
        
        doc.add_heading("4. VALOR GLOBAL", level=2)
        doc.add_paragraph(valor_global or "Não especificado")
        
        doc.add_heading("5. REAJUSTE", level=2)
        doc.add_paragraph(reajuste or "Não especificado")
        
        doc.add_heading("6. GARANTIAS", level=2)
        doc.add_paragraph(garantias or "Não especificado")
        
        doc.add_heading("7. PRAZOS E FORMA DE PAGAMENTO", level=2)
        doc.add_paragraph(prazos_pagamento or "Não especificado")
        
        doc.add_heading("8. OBRIGAÇÕES DA CONTRATADA", level=2)
        doc.add_paragraph(obrigacoes_contratada or "Não especificado")
        
        doc.add_heading("9. OBRIGAÇÕES DA CONTRATANTE", level=2)
        doc.add_paragraph(obrigacoes_contratante or "Não especificado")
        
        doc.add_heading("10. FISCALIZAÇÃO E ACOMPANHAMENTO", level=2)
        doc.add_paragraph(fiscalizacao or "Não especificado")
        
        doc.add_heading("11. PENALIDADES", level=2)
        doc.add_paragraph(penalidades or "Não especificado")
        
        doc.add_heading("12. RESCISÃO CONTRATUAL", level=2)
        doc.add_paragraph(rescisao or "Não especificado")
        
        doc.add_heading("13. FORO COMPETENTE", level=2)
        doc.add_paragraph(foro or "Não especificado")
        
        if observacoes_finais:
            doc.add_heading("14. OBSERVAÇÕES FINAIS", level=2)
            doc.add_paragraph(observacoes_finais)
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        st.download_button(
            label="📥 Baixar Minuta em DOCX",
            data=buffer,
            file_name=f"Contrato_TJSP_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        st.success("✅ Documento DOCX gerado com sucesso!")

st.caption("📎 Este módulo utiliza o backend refatorado (utils/integration_contrato.py) com lazy loading e modo degradado.")

