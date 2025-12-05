import sys
from pathlib import Path
ROOT = Path(__file__).resolve().parents[2]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

# ==============================
# pages/05_📑 TR – Termo de Referência.py
# SynapseNext – Secretaria de Administração e Abastecimento (TJSP)
# ==============================

import streamlit as st
from datetime import datetime
import os, sys, json
from io import BytesIO
from docx import Document
from utils.ui_components import aplicar_estilo_global, exibir_cabecalho_padrao
from utils.integration_tr import export_tr_to_json, ler_modelos_tr
from openai import OpenAI

# ==========================================================
# ⚙️ Configuração
# ==========================================================
st.set_page_config(page_title="📑 Termo de Referência", layout="wide", page_icon="📑")
aplicar_estilo_global()

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY") or os.getenv("openai_api_key")
client = OpenAI(api_key=OPENAI_API_KEY)

# ==========================================================
# 🏛️ Cabeçalho institucional
# ==========================================================
exibir_cabecalho_padrao(
    "📑 Termo de Referência (TR)",
    "Pré-preenchimento automático a partir de insumos + geração IA institucional"
)
st.divider()

# ==========================================================
# 🔍 Detecção e carregamento de insumos automáticos (com fallback persistente)
# ==========================================================
defaults = {}
EXPORTS_JSON_DIR = os.path.join("exports", "insumos", "json")

# Sessão ativa
if "tr_campos_ai" in st.session_state:
    defaults = st.session_state.get("tr_campos_ai", {})
    st.success("📎 Dados recebidos automaticamente do módulo INSUMOS (via sessão ativa).")

# Fallback: último insumo persistido
elif os.path.exists(EXPORTS_JSON_DIR):
    try:
        arquivos = sorted([f for f in os.listdir(EXPORTS_JSON_DIR) if f.endswith(".json")], reverse=True)
        if arquivos:
            caminho = os.path.join(EXPORTS_JSON_DIR, arquivos[0])
            with open(caminho, "r", encoding="utf-8") as f:
                dados = json.load(f)
            campos = dados.get("campos_ai", {})
            if isinstance(campos, dict):
                defaults = campos
                artefato = dados.get("artefato", "—")
                st.info(f"📎 Último insumo {artefato} carregado automaticamente ({arquivos[0]}).")
    except Exception as e:
        st.warning(f"⚠️ Falha ao recuperar insumo persistido: {e}")

# Nenhum insumo detectado
if not defaults:
    st.info("Nenhum insumo ativo detectado. Você pode preencher manualmente ou aguardar integração via módulo **🔧 Insumos**.")

# ==========================================================
# 🧾 Formulário TR – Estrutura institucional
# ==========================================================
st.subheader("📘 Entrada – Termo de Referência")

col1, col2 = st.columns(2)
with col1:
    objeto = st.text_area("Objeto da contratação", value=defaults.get("objeto", ""), height=120)
    justificativa_tecnica = st.text_area("Justificativa técnica", value=defaults.get("justificativa_tecnica", ""), height=120)
    especificacao_tecnica = st.text_area("Especificações técnicas", value=defaults.get("especificacao_tecnica", ""), height=120)
with col2:
    criterios_julgamento = st.text_area("Critérios de julgamento", value=defaults.get("criterios_julgamento", ""), height=120)
    riscos = st.text_area("Riscos associados", value=defaults.get("riscos", ""), height=120)
    observacoes_finais = st.text_area("Observações finais", value=defaults.get("observacoes_finais", ""), height=120)

st.divider()

col3, col4, col5 = st.columns(3)
with col3:
    prazo_execucao = st.text_input("Prazo de execução", value=defaults.get("prazo_execucao", ""))
with col4:
    estimativa_valor = st.text_input("Estimativa de valor (R$)", value=defaults.get("estimativa_valor", ""))
with col5:
    fonte_recurso = st.text_input("Fonte de recurso", value=defaults.get("fonte_recurso", ""))

# ==========================================================
# ⚙️ Botão de Processamento IA
# ==========================================================
st.divider()
st.subheader("⚙️ Geração de Artefato com IA Institucional")

if st.button("🤖 Gerar artefato com IA institucional"):
    with st.spinner("Gerando artefato completo com base nos dados e modelos do TJSP..."):
        modelos = ler_modelos_tr()
        campos = {
            "objeto": objeto,
            "justificativa_tecnica": justificativa_tecnica,
            "especificacao_tecnica": especificacao_tecnica,
            "criterios_julgamento": criterios_julgamento,
            "riscos": riscos,
            "observacoes_finais": observacoes_finais,
            "prazo_execucao": prazo_execucao,
            "estimativa_valor": estimativa_valor,
            "fonte_recurso": fonte_recurso,
        }

        user_prompt = f"""
Com base nos campos abaixo e nos modelos institucionais do TJSP, elabore o texto completo de um Termo de Referência (TR):

Campos:
{json.dumps(campos, ensure_ascii=False, indent=2)}

Modelos institucionais:
\"\"\"{modelos}\"\"\"


O texto deve seguir o padrão redacional e técnico do TJSP.
"""

        try:
            response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": "Você é um redator institucional do Tribunal de Justiça de São Paulo, responsável por elaborar termos de referência padronizados conforme as normas da SAAB/TJSP."},
                    {"role": "user", "content": user_prompt},
                ],
                temperature=0.3
            )
            artefato_tr = response.choices[0].message.content.strip()

            st.session_state["artefato_tr_gerado"] = artefato_tr
            st.success("✅ Artefato gerado com sucesso! Você pode agora exportá-lo como documento oficial (DOCX).")
            st.text_area("📄 Pré-visualização do artefato gerado:", artefato_tr, height=300)

        except Exception as e:
            st.error(f"Erro ao gerar artefato com IA: {e}")

# ==========================================================
# 💾 Exportação do artefato (DOCX)
# ==========================================================
if "artefato_tr_gerado" in st.session_state:
    artefato_tr = st.session_state["artefato_tr_gerado"]
    doc = Document()
    doc.add_heading("TERMO DE REFERÊNCIA", level=1)
    doc.add_paragraph(artefato_tr)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    st.download_button(
        label="📤 Exportar artefato em DOCX",
        data=buffer,
        file_name="TR_rascunho.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

st.caption("📎 O artefato acima é um rascunho institucional gerado pela IA com base nos modelos da SAAB/TJSP.")
