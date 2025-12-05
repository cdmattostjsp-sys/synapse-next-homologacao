import os
from docx import Document
from reportlab.pdfgen import canvas

BASE_PATH = os.path.join("tests", "data", "insumos_ficticios")
os.makedirs(BASE_PATH, exist_ok=True)

def create_docx(filename, title, sections):
    path = os.path.join(BASE_PATH, filename)
    doc = Document()
    doc.add_heading(title, 0)
    for section in sections:
        doc.add_heading(section["titulo"], level=1)
        doc.add_paragraph(section["conteudo"])
    doc.save(path)
    return path

def create_pdf(filename, title, content):
    path = os.path.join(BASE_PATH, filename)
    pdf = canvas.Canvas(path)
    pdf.setFont("Helvetica", 14)
    pdf.drawString(100, 780, title)
    pdf.setFont("Helvetica", 10)
    pdf.drawString(100, 750, content)
    pdf.save()
    return path

def gerar_insumos_ficticios():
    exemplos = [
        {
            "arquivo": "DFD_exemplo.docx",
            "titulo": "Diagrama de Fluxo de Dados – Exemplo Fictício",
            "secoes": [
                {"titulo": "Contexto", "conteudo": "Fluxo simulado entre setor de compras e SAAB."},
                {"titulo": "Objetivo", "conteudo": "Demonstrar integração sintética para teste do pipeline semântico."},
            ],
        },
        {
            "arquivo": "ETP_exemplo.docx",
            "titulo": "Estudo Técnico Preliminar – Exemplo Fictício",
            "secoes": [
                {"titulo": "Escopo", "conteudo": "Avaliar soluções técnicas para aquisição de sistemas integrados."},
                {"titulo": "Justificativa", "conteudo": "Documento criado para fins de homologação."},
            ],
        },
        {
            "arquivo": "TR_exemplo.docx",
            "titulo": "Termo de Referência – Exemplo Fictício",
            "secoes": [
                {"titulo": "Objeto", "conteudo": "Contratação de solução integrada de apoio à gestão de licitações."},
                {"titulo": "Especificações", "conteudo": "Simulação textual de parâmetros técnicos para validação."},
            ],
        },
    ]
    for ex in exemplos:
        create_docx(ex["arquivo"], ex["titulo"], ex["secoes"])

    create_pdf(
        "edital_exemplo.pdf",
        "Minuta de Edital Fictício",
        "Documento de exemplo para testes de parsing e extração semântica."
    )

if __name__ == "__main__":
    gerar_insumos_ficticios()
    print(f"✅ Insumos fictícios gerados em: {BASE_PATH}")
